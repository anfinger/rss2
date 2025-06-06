# -*- coding: utf8 -*-

from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from .forms import ReiseForm
from uuid import UUID
import re
import locale
from django.utils import timezone
from django import db
from collections import namedtuple, OrderedDict 
from django.db import connection
from itertools import chain
from django.core import serializers
from django.template.loader import render_to_string
from django.db.models import Min, F, Value as V, Q
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime
from datetime import timedelta
import time
from docx import *
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor
from StringIO import StringIO
from django.contrib.staticfiles.templatetags.staticfiles import static
import mysql.connector
import json
from HTMLParser import HTMLParser

from django.utils.encoding import python_2_unicode_compatible

#from django.core.files.storage import DefaultStorage
#from filebrowser.sites import FileBrowserSite
#from filebrowser.sites import site

#from django.core.files.storage import FileSystemStorage

from .models import Reise, Reisetermine, Reisezielregionen, Zielregion, Reisekategorien, Reisehinweise, Ausflugspakete, Ausflugspaketpreise, AusflugspaketeZuReisetagen, LeistungenAusflugspaket, Abfahrtszeiten, LeistungenReise, Reisebeschreibung, Reisetage, Reisepreise, Preis, ReisepreisZusatz, Zusatzleistung, Fruehbucherrabatt, Reisebilder, Reisekatalogzugehoerigkeit, Katalog

from .querysets import get_detail_hinweise_queryset, get_detail_angebote_queryset, get_detail_auftragsbestaetigungen_queryset

# Create your views here.

#site = FileBrowserSite(name='huhu', storage=DefaultStorage())
#site.directory = "/images/"
#site.storage.location = DefaultStorage().location

def namedtuplefetchall(cursor):
    "Return all rows from a cursor as a namedtuple"
    desc = cursor.description
    nt_result = namedtuple('Result', [col[0] for col in desc])
    return [nt_result(*row) for row in cursor.fetchall()]

################################################################t.META.HTTP_REFERER }}
# Index Seite, Reisen Übersicht                                  #
##################################################################
def index(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT reiseID, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) ORDER BY RT.min_datum;")
    #cursor.execute("SELECT reiseID, reisen_reise.untertitel, einleitung, reisetyp, katalogseite, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reiseID) LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id);");
    termine = namedtuplefetchall(cursor)
    cursor.close()
    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)
    return render(request, 'reisen/index.html', {'termine': termine, 'dibug': dibug})

##################################################################
# Winterreisen 2016 2017                                         #
##################################################################
def winter2016_17(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '5fcadf9947314b82ac79802e92585c39' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    cursor = connection.cursor()
    cursor.execute("select DISTINCT reisen_reise.titel, korrektur_bemerkung_intern from reisen_reise left join reisen_reisekatalogzugehoerigkeit on (reiseID = reise_id_id) where katalog_id_id = '5fcadf9947314b82ac79802e92585c39';")
    korrekturen = namedtuplefetchall(cursor)
    cursor.close()

    return render(request, 'reisen/index.html', {'termine': termine, 'dibug': dibug, 'korrekturen': korrekturen })

##################################################################
# Winterreisen 2016 2017                                         #
##################################################################
def winter2017_18(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, neu, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '21a8a8c913854f41865953f6f10f538f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    korrekturen = Reise.objects.select_related().filter(
        reisekatalogzugehoerigkeit__katalog_id='21a8a8c913854f41865953f6f10f538f'
      ).all()#.values(
#        'reiseID', 
 #       'titel',
  #      'korrektur_bemerkung_intern'
   #   ).distinct()

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    #cursor = connection.cursor()
    #cursor.execute("select DISTINCT reisen_reise.titel, korrektur_bemerkung_intern from reisen_reise left join reisen_reisekatalogzugehoerigkeit on (reiseID = reise_id_id) where katalog_id_id = '21a8a8c913854f41865953f6f10f538f';")
    #korrekturen = namedtuplefetchall(cursor)
    #cursor.close()

    return render(request, 'reisen/index.html', {'termine': termine, 'dibug': dibug, 'korrekturen': korrekturen })

##################################################################
# Winterreisen 2018 2019                                         #
##################################################################
def winter2018_19(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, neu, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '278d87913b644f658dcf6af9c734843d' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    korrekturen = Reise.objects.select_related().filter(
        reisekatalogzugehoerigkeit__katalog_id='278d87913b644f658dcf6af9c734843d'
      ).all()#.values(
#        'reiseID',
 #       'titel',
  #      'korrektur_bemerkung_intern'
   #   ).distinct()

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    #cursor = connection.cursor()
    #cursor.execute("select DISTINCT reisen_reise.titel, korrektur_bemerkung_intern from reisen_reise left join reisen_reisekatalogzugehoerigkeit on (reiseID = reise_id_id) where katalog_id_id = '21a8a8c913854f41865953f6f10f538f';")
    #korrekturen = namedtuplefetchall(cursor)
    #cursor.close()

    return render(request, 'reisen/index.html', {'termine': termine, 'dibug': dibug, 'korrekturen': korrekturen })


##################################################################
# Musicalfahrten                                                 #
##################################################################
#def musicalfahrten(request):
#
#    termine = Reisetermine.objects.select_related().filter(
#        datum_beginn__gte=datetime.now(),
#        reise_id__status='f'
#      ).filter(reise_id__reisetyp='Musicalfahrt')
#      ).order_by('datum_beginn').values(
#        'datum_beginn',
#        'markierung',
#        'kommentar',
#        'reise_id__reiseID',
#        'reise_id__titel',
#        'reise_id__sonstigeReisebeschreibung_titel'
#      ).distinct()
#
#    musicalfahrten = [
#      {
#        "monat": (int(datetime.strftime(termin['datum_beginn'], "%-m"))-int(datetime.strftime(termine[i-1]['datum_beginn'], "%-m")) if i>0 else 1),
#        "datum_beginn": termin['datum_beginn'],
#        "datum_kommentar": termin['kommentar'],
#        "datum_markierung": termin['markierung'],
#        "reisedaten": get_object_or_404(Reise, pk=termin['reise_id__reiseID']),
#        "abfahrtszeiten": Abfahrtszeiten.objects.filter(reise_id=termin['reise_id__reiseID']).order_by('position'),
#        "preise": [
#          {
#            "preistitel": str(preis.preis_id),
#            "preis": preis.preis,
#            "markierung": preis.markierung,
#            "kommentar": preis.kommentar,
#            "zpreise": [
#              {
#                "zpreistitel": str(zpreis.preis_id),
#                "zpreis": zpreis.preis
#              }
#              for zpreis in ReisepreisZusatz.objects.filter(reisepreis_id=preis.reisepreisID)
#            ]
#          }
#          for preis in Reisepreise.objects.filter(reise_id=termin['reise_id__reiseID'])
#        ],
#        "hinweise": [
#          {
#            "hinweis": str(hinweis.hinweis_id)
#          }
#          for hinweis in Reisehinweise.objects.filter(reise_id=termin['reise_id__reiseID'])
#        ],
#      }
#      for i, termin in enumerate(termine)
#    ]
#
#    ausgabeformat = request.GET.get('format')
#
#    return render(request, 'reisen/musicals.html', {'musicalfahrten': musicalfahrten })


##################################################################
# Tagesfahrten                                                   #
##################################################################
def tagesfahrten(request):

    termine = Reisetermine.objects.select_related().filter(
        #if ausgabeformat == 'xml':
        #  datum_beginn__gtedatum_beginn__gte=datetime.datetime(2018, 11, 30, 00, 00),
        #else:
        #  datum_beginn__gte=datetime.now(),
        datum_beginn__gte=datetime.now(),
        reise_id__status='f'
      ).filter(
        Q(reise_id__reisetyp='Tagesfahrt') | Q(reise_id__reisetyp='Musicalfahrt')
      ).order_by('datum_beginn').values(
        'datum_beginn',
        'markierung',
        'kommentar',
        'reise_id__reiseID',
        'reise_id__titel',
        'reise_id__sonstigeReisebeschreibung_titel'
      ).distinct()

    tagesfahrten = [
      {
        "monat": (int(datetime.strftime(termin['datum_beginn'], "%-m"))-int(datetime.strftime(termine[i-1]['datum_beginn'], "%-m")) if i>0 else 1),
        "datum_beginn": termin['datum_beginn'],
        "datum_kommentar": termin['kommentar'],
        "datum_markierung": termin['markierung'],
        "reisedaten": get_object_or_404(Reise, pk=termin['reise_id__reiseID']),
        "abfahrtszeiten": Abfahrtszeiten.objects.filter(reise_id=termin['reise_id__reiseID']).order_by('position'),
        "preise": [
          {
            "preistitel": str(preis.preis_id),
            "preis": preis.preis,
            "markierung": preis.markierung,
            "kommentar": preis.kommentar,
            "zpreise": [
              {
                "zpreistitel": str(zpreis.preis_id),
                "zpreis": zpreis.preis
              }
              for zpreis in ReisepreisZusatz.objects.filter(reisepreis_id=preis.reisepreisID)
            ]
          }
          for preis in Reisepreise.objects.filter(reise_id=termin['reise_id__reiseID'])
        ],
        "hinweise": [
          {
            "hinweis": str(hinweis.hinweis_id)
          }
          for hinweis in Reisehinweise.objects.filter(reise_id=termin['reise_id__reiseID'])
        ],
      }
      for i, termin in enumerate(termine)
    ]

    ausgabeformat = request.GET.get('format')

    if ausgabeformat == 'xml':

      return render(request, 'reisen/tagesfahrten.xml', {'tagesfahrten': tagesfahrten })

    elif ausgabeformat != 'docx':

      return render(request, 'reisen/tagesfahrten.html', {'tagesfahrten': tagesfahrten })

    else:

      locale.setlocale(locale.LC_TIME, "de_DE")

      document = Document()
      docx_title="tagesfahrten_" + datetime.strftime(datetime.now(), "%m") + ".docx"
      #document.add_picture('/var/www/reiseservice-schwerin/rss2' + static('images/logo_google_blau.png'), width=Inches(2))
      #sections = document.sections
      #sections[0].orientation = WD_ORIENT.LANDSCAPE
      #document.sections[0].orientation = WD_ORIENT.LANDSCAPE
      font = document.styles['Normal'].font
      font.name = 'Idea'
      font.size = Pt(72)
      document.add_paragraph()
      p = document.add_paragraph()
      #p.add_run("%s" % datetime.strftime(datetime.now(), "%B")).bold = True
      p.add_run(u'Tagesfahrten Monat').bold = True
      document.add_page_break()
      font.name = 'Submariner R24'
      font.size = Pt(12)
      paragraph_format = document.styles['Normal'].paragraph_format
      paragraph_format.space_before = Pt(0)
      paragraph_format.space_after = Pt(0)
      paragraph_format.line_spacing = 1

      orte = {
        'HBF': u'Hbf. Schwerin',
        'VSB': u'Hast. v. Stauffenberg Str.',
        'GAR': u'Gartenstadt',
        'WIS': u'ZOB Wismar',
        'ROG': 'Gadebusch Roggendorfer Str.',
        'ANK': u''
      }

      for tagesfahrt in tagesfahrten:
        p = document.add_paragraph()
        p.add_run(datetime.strftime(tagesfahrt['datum_beginn'], "%a., %d.%m.") + "\t" + tagesfahrt['reisedaten'].titel).bold = True
        if tagesfahrt['reisedaten'].untertitel or tagesfahrt['reisedaten'].einleitung:
          p = document.add_paragraph()
        if tagesfahrt['reisedaten'].untertitel:
          p.add_run(tagesfahrt['reisedaten'].untertitel + ' ')
        if  tagesfahrt['reisedaten'].einleitung:
          p.add_run(tagesfahrt['reisedaten'].einleitung)
        if tagesfahrt['reisedaten'].veranstalter != 'RS':
          if tagesfahrt['reisedaten'].veranstalter == 'SH':
            p = document.add_paragraph()
            p.add_run('Veranstalter: Sewert Reisen').italic = True
        if tagesfahrt['hinweise']:
          for hinweis in tagesfahrt['hinweise']:
            p = document.add_paragraph()
            p.add_run(unicode(hinweis['hinweis'], "utf-8")).italic = True
        for idx, abfahrtszeit in enumerate(tagesfahrt['abfahrtszeiten']):
          if idx == 0:
            p = document.add_paragraph()
            p.add_run('Abfahrt: ').name = 'Submariner R24 light'
          if idx == (len(tagesfahrt['abfahrtszeiten'])-1):
            p = document.add_paragraph()
            p.add_run('Ankunft: ').name = 'Submariner R24 light'
          p.add_run(str(abfahrtszeit.zeit)[:5] + ' Uhr ' + orte[abfahrtszeit.ort]).name = 'Submariner R24 light'
          if idx < (len(tagesfahrt['abfahrtszeiten'])-2):
            p.add_run(', ').name = 'Submariner R24 light'

        newparagraph = False
        for idx, preis in enumerate(tagesfahrt['preise']):
          if preis['markierung']:
            if preis['markierung'] == tagesfahrt['datum_markierung']:
              if newparagraph:
                p = document.add_paragraph()
              newparagraph = True
              p.add_run('\t\t' + unicode(preis['preistitel'], "utf-8") + ': ').bold = True
              if preis['kommentar']:
                p.add_run(preis['kommentar'] + ' ').bold = True
              p.add_run(str(preis['preis']) + u' €').bold = True
          else:
            if newparagraph:
              p = document.add_paragraph()
            newparagraph = True
            p.add_run('\t\t' + unicode(preis['preistitel'], "utf-8") + ': ').bold = True
            if preis['kommentar']:
              p.add_run(preis['kommentar'] + ' ').bold = True
            p.add_run(str(preis['preis']) + u' €').bold = True
            
        p = document.add_paragraph()
        p.add_run().size = Pt(8)
        
      # Prepare document for download        
      # -----------------------------
      f = StringIO()
      document.save(f)
      length = f.tell()
      f.seek(0)
      response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
      )
      response['Content-Disposition'] = 'attachment; filename=' + docx_title
      response['Content-Length'] = length
      return response 
      

##################################################################
# XML Export Winterreisen 2016 2017                              #
##################################################################
def winter2016_17_export(request):

    #XMLSerializer = serializers.get_serializer("xml")
    #xml_serializer = XMLSerializer()
    #xml_serializer.serialize(Reise.objects.all())
    #data = xml_serializer.getvalue()

    #with open("/var/www/reiseservice-schwerin/rss2/media/termine.xml", "w") as out:
    #    xml_serializer.serialize(Reise.objects.all(), stream=out)

    #xml = render_to_string('xml_template.xml', {'query_set': query_set})

    cursor = connection.cursor()
    cursor.execute("SELECT reiseID, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '5fcadf9947314b82ac79802e92585c39' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    return render(request, 'reisen/export.xml', {'termine': termine })

##################################################################
# Winterreisen 2016 2017                                         #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def winter1617(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '5fcadf9947314b82ac79802e92585c39' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- und Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- und Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2016 2017                                         #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def winter1718(request):
    dibug = ''
    cursor = connection.cursor()
    #cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '21a8a8c913854f41865953f6f10f538f' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '278d87913b644f658dcf6af9c734843d' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)
    cursor.close()   

    kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    #return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

    reiseziel = [u'Deutschland',u'Frankreich',u'Italien',u'Polen',u'Japan',u'Portugal']
    reisekat = [u'kombinierte Flug- & Busreisen',u'Busreisen',u'Kuren, Gesundheits- und Wellnessreisen']

    reiseids = Reise.objects.select_related().filter(
        reisetermine__datum_beginn__gte=datetime.now(),
        status='f',
      ).order_by('reiseID').values(
        'reiseID',
        'titel',
        'sonstigeReisebeschreibung_titel'
    ).distinct()

#    reise = get_object_or_404(Reise, pk=pk)
#    termine = Reisetermine.objects.filter(reise_id=pk).order_by('datum_beginn')
#    abfahrtszeiten = Abfahrtszeiten.objects.filter(reise_id=pk).order_by('position')
#    leistungen = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=0).order_by('position')
#    leistungennichtindividual = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=1).order_by('position')
#    zusatzleistungen = Zusatzleistung.objects.filter(reise_id=pk).order_by('position')
#    fruehbucherrabatte = Fruehbucherrabatt.objects.filter(reise_id=pk).order_by('datum_bis')
#    reisebeschreibung = Reisebeschreibung.objects.filter(reise_id=pk).order_by('position')

#    for idx, tag in enumerate(tage):
#      tag.reisetagID = str(tag.reisetagID).replace('-','')
#      if tag.tagnummertext:
#        tag.nummerntext = tag.tagnummertext
#      else:
#        naechster_tag = tage[(idx+1) % len(tage)]
#        if len(tage) == 1:
#            tag.nummerntext = ''
#        elif (naechster_tag.tagnummer == (tag.tagnummer + 1)) or (idx == (len(tage)-1)):
#            tag.nummerntext = str(tag.tagnummer) + '. Tag:'
#        else:
#            tag.nummerntext = str(tag.tagnummer) + '. - ' + str(naechster_tag.tagnummer-1) + '. Tag:'

    nested_reisen = [
      {
        "reiseID": rid['reiseID'],
        "titel": rid['titel'],
        "sonstigeReisebeschreibung_titel": rid['sonstigeReisebeschreibung_titel'],
	#"reisedaten": Reise.objects.filter(reiseID=rid['reiseID']),
	"reisedaten": get_object_or_404(Reise, pk=rid['reiseID']),
        "termine": [
          {
            "beginn": termin.datum_beginn,
            "ende": termin.datum_ende,
            "markierung": termin.markierung
          }
          for termin in Reisetermine.objects.filter(reise_id=rid['reiseID'],datum_beginn__gte=datetime.now())
        ],
        "tage": Reisetage.objects.filter(reise_id=rid['reiseID']).order_by('tagnummer'),
        "leistungen": LeistungenReise.objects.filter(reise_id=rid['reiseID']).filter(nichtindividual=0).order_by('position'),
        "leistungennichtindividual": LeistungenReise.objects.filter(reise_id=rid['reiseID']).filter(nichtindividual=1).order_by('position'),
        "abfahrtszeiten": Abfahrtszeiten.objects.filter(reise_id=rid['reiseID']).order_by('zeit'),
        "reisebeschreibungen": Reisebeschreibung.objects.filter(reise_id=rid['reiseID']).order_by('position'),
        "preise": [
          {
            "preistitel": str(preis.preis_id),
            "preis": preis.preis,
            "markierung": preis.markierung,
            "kommentar": preis.kommentar,
            "zpreise": [
              {
                "zpreistitel": str(zpreis.preis_id),
                "zpreis": zpreis.preis
              }
              for zpreis in ReisepreisZusatz.objects.filter(reisepreis_id=preis.reisepreisID)
            ]
          }
          for preis in Reisepreise.objects.filter(reise_id=rid['reiseID'])
        ],
        "zielregionen": [
          {
            "zielregion": str(zielregion.zielregion_id)
          }
          for zielregion in Reisezielregionen.objects.filter(reise_id=rid['reiseID'])
        ],
        "kategorien": [
          {
            "kategorie": str(kategorie.kategorie_id)
          }
          for kategorie in Reisekategorien.objects.filter(reise_id=rid['reiseID'])
        ],
        "zusatzleistungen": [
          {
            "zltitel": zl.titel,
            "zlkommentar": zl.kommentar,
            "zlkommentar_titel": zl.kommentar_titel,
            "zlpreise": [
              {
                "zlpreis": zlpreis.preis 
              }
              for zlpreis in Ausflugspaketpreise.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ],
            "apleistungen": [
              {
                "apleistung": apleistung.leistung
              }
              for apleistung in LeistungenAusflugspaket.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ],
            "aptage": [
              {
                "aptag": str(aptag.reisetag_id)
              }
              for aptag in AusflugspaketeZuReisetagen.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ]
          }
          for zl in Ausflugspakete.objects.filter(reise_id=rid['reiseID'])
        ],
        "hinweise": [
          {
            "hinweis": str(hinweis.hinweis_id)
          }
          for hinweis in Reisehinweise.objects.filter(reise_id=rid['reiseID'])
        ]
      }
      for rid in reiseids
    ]
    #nested_reisen = [{"reiseID": reise.reiseID, "termine": [{"beginn": termin.datum_beginn, "ende": termin.datum_ende} for termin in reise.reisetermine__reiseID.all()]} for reise in reisen]
    #nested_reisen = [{"reiseID": reise['reiseID'], "termine": [{"beginn": termin.datum_beginn, "ende": termin.datum_ende} for termin in reise['reisetermine']]} for reise in reisen]

    ausgabeformat = request.GET.get('format')
    if ausgabeformat != 'xml':
      return render(request, 'reisen/index_reisen_web_alt.html', { 'termine': termine, 'dibug': dibug, 'kategorien': kategorien })
    else:
      return render(request, 'reisen/rb24_reiseservice-schwerin.xml', { 'reiseids': reiseids, 'reisen': list(nested_reisen), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien })


##################################################################
# Reisebus24 Export                                              #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def rb24(request):

    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '21a8a8c913854f41865953f6f10f538f' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    #return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

    reiseziel = [u'Deutschland',u'Frankreich',u'Italien',u'Polen',u'Japan',u'Portugal']
    reisekat = [u'kombinierte Flug- & Busreisen',u'Busreisen',u'Kuren, Gesundheits- und Wellnessreisen']

    reiseids = Reise.objects.select_related().filter(
        reisetermine__datum_beginn__gte=datetime.now(),
        status='f',
      ).order_by('reiseID').values(
        'reiseID',
        'titel',
        'sonstigeReisebeschreibung_titel'
    ).distinct()

    nested_reisen = [
      {
        "reiseID": rid['reiseID'],
        "titel": rid['titel'],
        "sonstigeReisebeschreibung_titel": rid['sonstigeReisebeschreibung_titel'],
        #"reisedaten": Reise.objects.filter(reiseID=rid['reiseID']),
        "reisedaten": get_object_or_404(Reise, pk=rid['reiseID']),
        "termine": [
          {
            "beginn": termin.datum_beginn,
            "ende": termin.datum_ende,
            "markierung": termin.markierung
          }
          for termin in Reisetermine.objects.filter(reise_id=rid['reiseID'],datum_beginn__gte=datetime.now())
        ],
        "tage": Reisetage.objects.filter(reise_id=rid['reiseID']).order_by('tagnummer'),
        "leistungen": LeistungenReise.objects.filter(reise_id=rid['reiseID']).filter(nichtindividual=0).order_by('position'),
        "leistungennichtindividual": LeistungenReise.objects.filter(reise_id=rid['reiseID']).filter(nichtindividual=1).order_by('position'),
        "abfahrtszeiten": Abfahrtszeiten.objects.filter(reise_id=rid['reiseID']).order_by('zeit'),
        "reisebeschreibungen": Reisebeschreibung.objects.filter(reise_id=rid['reiseID']).order_by('position'),
        "preise": [
          {
            "preistitel": str(preis.preis_id),
            "preis": preis.preis,
            "markierung": preis.markierung,
            "kommentar": preis.kommentar,
            "zpreise": [
              {
                "zpreistitel": str(zpreis.preis_id),
                "zpreis": zpreis.preis
              }
              for zpreis in ReisepreisZusatz.objects.filter(reisepreis_id=preis.reisepreisID)
            ]
          }
          for preis in Reisepreise.objects.filter(reise_id=rid['reiseID'])
        ],
        "zielregionen": [
          {
            "zielregion": str(zielregion.zielregion_id)
          }
          for zielregion in Reisezielregionen.objects.filter(reise_id=rid['reiseID'])
        ],
        "kategorien": [
          {
            "kategorie": str(kategorie.kategorie_id)
          }
          for kategorie in Reisekategorien.objects.filter(reise_id=rid['reiseID'])
        ],
        "zusatzleistungen": [
          {
            "zltitel": zl.titel,
            "zlkommentar": zl.kommentar,
            "zlkommentar_titel": zl.kommentar_titel,
            "zlpreise": [
              {
                "zlpreis": zlpreis.preis
              }
              for zlpreis in Ausflugspaketpreise.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ],
            "apleistungen": [
              {
                "apleistung": apleistung.leistung
              }
              for apleistung in LeistungenAusflugspaket.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ],
            "aptage": [
              {
                "aptag": str(aptag.reisetag_id)
              }
              for aptag in AusflugspaketeZuReisetagen.objects.filter(ausflugspaket_id=zl.ausflugspaketID)
            ]
          }
          for zl in Ausflugspakete.objects.filter(reise_id=rid['reiseID'])
        ],
        "hinweise": [
          {
            "hinweis": str(hinweis.hinweis_id)
          }
          for hinweis in Reisehinweise.objects.filter(reise_id=rid['reiseID'])
        ]
      }
      for rid in reiseids
    ]
    
    return render(request, 'reisen/rb24_reiseservice-schwerin.xml', { 'reiseids': reiseids, 'reisen': list(nested_reisen), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2018 2019                                         #
##################################################################
def winter1920(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '1f9e043d4f4341a8a563d408fd267c56' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2020 2021                                         #
##################################################################
def winter2021(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '3208977df8c640828e81276cead39a1d' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2021 2022                                         #
##################################################################
def winter2122(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '07451f76a08c43c0ba3f44cfa8579475' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2022 2023                                         #
##################################################################
def winter2223(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '49e741e8ba0c44dd9ecd663b9854cbd5' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2023 2024                                         #
##################################################################
def winter2324(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '1359062932574bb888958c5b24f69060' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2024 2025                                         #
##################################################################
def winter2425(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '7ea72a2dd81d4c55b1d06231afc30f3f' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

##################################################################
# Winterreisen 2025 2026                                         #
##################################################################
def winter2526(request):

    dibug = ''

    cursor = connection.cursor()

    cursor.execute("SELECT DISTINCT reiseID, RP.abpreis, reisen_kategorie.kategorie, reisen_reise.untertitel, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, reisen_reise.sonstigeReisebeschreibung_titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn, datum_ende ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '91047b6622874816af2ee870908018c7' AND reisen_kategorie.kategorie in ('Musicals & Shows','Weihnachts- & Silvesterreisen', 'Adventsreisen & Weihnachtsmärkte', 'Kuren, Gesundheits- und Wellnessreisen', 'Flusskreuzfahrten', 'Ostern', 'kombinierte Flug- & Busreisen', 'Frühlingsreisen', 'Herbstreisen', 'Winterreisen') ORDER BY RT.min_datum;")

    termine = namedtuplefetchall(cursor)

    cursor.close()

    #kategorien = ['Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', 'Ostern', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']
    kategorien = [u'Herbstreisen', u'Adventsreisen & Weihnachtsmärkte', u'Weihnachts- & Silvesterreisen', 'Winterreisen', u'Frühlingsreisen', u'Kuren, Gesundheits- und Wellnessreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Musicals & Shows']

    #dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    return render(request, 'reisen/index_reisen_web_alt.html', {'termine': termine, 'dibug': dibug, 'kategorien': kategorien })

########################################################
# XML Export alle Reisen nach Termin Winter 2025/26     #
########################################################
def reiseuebersichtwinter(request):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")
    cursor.execute("""
      SELECT
        reiseID,
        reisen_reise.korrektur_bemerkung_intern AS intern,
        kat.groupkat AS k,
        reisen_reise.individualbuchbar,
        reisen_reise.neu,
        reisen_reise.titel AS Reiseziel,
        reisen_reise.sonstigeReisebeschreibung_titel AS Zusatztitel,
        CONCAT(
          DATE_FORMAT(reisen_reisetermine.datum_beginn,'%d.%m.'),
          '-',
          DATE_FORMAT(reisen_reisetermine.datum_ende,'%d.%m.%y')
        ) AS Termin,
        (TO_DAYS(reisen_reisetermine.datum_ende)-TO_DAYS(reisen_reisetermine.datum_beginn)+1) AS Tage,
        CONCAT(
          katalogseite,
          IF(anzahl_seiten_im_katalog > 1, concat('|',katalogseite+1), '')
        ) AS Seite
        FROM
          reisen_reise
        LEFT JOIN
          reisen_reisetermine
        ON
          reisen_reise.reiseID = reisen_reisetermine.reise_id_id
        LEFT JOIN
          (
            SELECT
              reise_id_id,
              GROUP_CONCAT(kategorie) AS groupkat
            FROM
              reisen_reisekategorien
            LEFT JOIN
              reisen_kategorie
            ON
              kategorieID = kategorie_id_id
            WHERE
              kategorie in ('kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Busreisen')
            GROUP BY
              reise_id_id
          ) AS kat
        ON
          kat.reise_id_id = reisen_reise.reiseID
        LEFT JOIN
          reisen_reisekatalogzugehoerigkeit
        ON
          reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id
        WHERE
          reisen_reisekatalogzugehoerigkeit.katalog_id_id = '91047b6622874816af2ee870908018c7'
        AND
          katalogseite > 0
        ORDER BY
          FIELD(
            k,
            'Busreisen',
            'kombinierte Flug- & Busreisen',
            'Flusskreuzfahrten'
          ),
          katalogseite,
          datum_beginn,
          datum_ende,
          position_auf_seite;
        """)

    termine = namedtuplefetchall(cursor)

    swine = False
    skifahren = False
    gruenewoche = False
    barcelona = False
    danke = False
    blau = False
    azoren = False
    rom = False
    andalusien = False
    lissabon = False
    starlight = False
    elfi = False
    wellnesswochenende = False
    kategorie = ''
    for i in range(len(termine)):
      if termine[i].k != kategorie:
        kategorie = termine[i].k
      else:
        termine[i] = termine[i]._replace(k = '')
      if u'Weihnachten und Silvester an der polnischen Ostseeküste' in termine[i-1].Reiseziel and u'Weihnachten und Silvester an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Blauf' in termine[i].Reiseziel and blau != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25')
        blau = True
      elif u'Blauf' in termine[i].Reiseziel and blau != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Starlight' in termine[i].Reiseziel and starlight != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & 26')
        starlight = True
      elif u'Starlight' in termine[i].Reiseziel and starlight != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Elbphil' in termine[i].Reiseziel and elfi != True:
        termine[i] = termine[i]._replace(Termin = 'Dez. 25 & Apr. 26')
        elfi = True
      elif u'Elbphil' in termine[i].Reiseziel and elfi != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Wellnesswochenende' in termine[i].Reiseziel and wellnesswochenende != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        wellnesswochenende = True
      elif u'Wellnesswochenende' in termine[i].Reiseziel and wellnesswochenende != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Azoren' in termine[i].Reiseziel and azoren != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & Apr. 26')
        azoren = True
      elif u'Azoren' in termine[i].Reiseziel and azoren != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'ewige Stadt' in termine[i].Reiseziel and rom != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & 26')
        rom = True
      elif u'ewige Stadt' in termine[i].Reiseziel and rom != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Andalusien' in termine[i].Reiseziel and andalusien != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        andalusien = True
      elif u'Andalusien' in termine[i].Reiseziel and andalusien != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Dankesch' in termine[i].Reiseziel and danke != True:
        termine[i] = termine[i]._replace(Termin = 'Nov. 25 & Jan. 26')
        danke = True
      elif u'Dankesch' in termine[i].Reiseziel and danke != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Swinemünde' in termine[i].Reiseziel and swine != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 bis März 26')
        termine[i] = termine[i]._replace(Tage = '8')
        swine = True
      elif u'Swinemünde' in termine[i].Reiseziel and swine != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Barcelona' in termine[i].Reiseziel and barcelona != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        barcelona = True
      elif u'Barcelona' in termine[i].Reiseziel and barcelona != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Lissabon' in termine[i].Reiseziel and lissabon != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & Mai 26')
        lissabon = True
      elif u'Lissabon' in termine[i].Reiseziel and lissabon != False:
        termine[i] = termine[i]._replace(Termin = '')
      #if u'Skifahren in Südtirol' in termine[i].Reiseziel and skifahren != True:
      #  termine[i] = termine[i]._replace(Termin = 'Jan. & Feb. 24')
      #  skifahren = True
      #elif u'Skifahren in Südtirol' in termine[i].Reiseziel and skifahren != False:
      #  termine[i] = termine[i]._replace(Termin = '')
      if not termine[i].Tage:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Grünen Woche' in termine[i].Reiseziel and gruenewoche == False:
        termine[i] = termine[i]._replace(Termin = '22. bis 27.01.25')
        termine[i] = termine[i]._replace(Tage = '1')
        gruenewoche = True
      elif u'Grünen Woche' in termine[i].Reiseziel and gruenewoche == True:
        termine[i] = termine[i]._replace(Termin = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reiseziel = termine[i].Reiseziel + ' (auch individuell buchbar)')
      #if (termine[i].anzahl_seiten_im_katalog == 0) and (termine[i].position_auf_seite == 1) and termine[i-1].Reiseziel != '':
        #termine[i-1] = termine[i-1]._replace(Reiseziel = termine[i-1].Reiseziel + ' (auch individuell buchbar)')
        #termine[i] = termine[i]._replace(Termin = '')

    cursor.close()

    ausgabeformat = request.GET.get('format')
    if ausgabeformat == 'html':
      return render(request, 'reisen/export_reiseuebersichtwinter.html', {'termine': termine })
    else:
      return render(request, 'reisen/export_reiseuebersichtwinter.xml', {'termine': termine })

########################################################
# XML Export alle Reisen nach Termin Sommer  2017      #
########################################################
def reiseterminuebersicht(request):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")
    cursor.execute("""
                    SELECT
                      reisen_reise.reiseID AS pk,
                      reisen_reise.korrektur_bemerkung_intern AS intern,
                      Veranstalter,
                      DAYOFYEAR(datum_beginn) AS tmp,
                      0 AS tmp2,
                      0 AS bus,
                      DATE_FORMAT(reisen_reisetermine.datum_beginn, '%M') AS Monat,
                      reisen_reise.individualbuchbar,
                      reisen_reise.neu,
                      reisen_reise.titel AS Reiseziel,
                      reisen_reisetermine.datum_beginn,
                      reisen_reisetermine.datum_ende,
                      CONCAT(DATE_FORMAT(reisen_reisetermine.datum_beginn,'%d.%m.'), '-', DATE_FORMAT(reisen_reisetermine.datum_ende,'%d.%m.%y')) AS Termin,
                      (TO_DAYS(reisen_reisetermine.datum_ende)-TO_DAYS(reisen_reisetermine.datum_beginn)+1) AS Tage,
                      katalogseite,
                      anzahl_seiten_im_katalog,
                      position_auf_seite,
                      CONCAT(katalogseite,IF(anzahl_seiten_im_katalog>1,CONCAT('|',katalogseite+1),'')) AS Seite
                    FROM
                      reisen_reise
                    LEFT JOIN
                      reisen_reisetermine
                    ON
                      (reisen_reise.reiseID = reisen_reisetermine.reise_id_id)
                    LEFT JOIN
                      reisen_reisekatalogzugehoerigkeit
                    ON
                      (reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id)
                    WHERE
                      reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2'
                    AND
                      DATE_FORMAT(reisen_reisetermine.datum_beginn, '%Y') > 2024
                    ORDER BY
                      datum_beginn,
                      datum_ende,
                      katalogseite,
                      position_auf_seite;
                  """)
    
    termine = namedtuplefetchall(cursor)

    month = ''
    dates = []
    delta = 0
    for i in range(len(termine)):
      if i == 0:
        dates.append(termine[i].datum_ende)
        termine[i] = termine[i]._replace(tmp = termine[i].tmp-1)
        if termine[i].Veranstalter == 'RS':
          termine[i] = termine[i]._replace(bus = 1)
      if termine[i].Monat != month:
        month = termine[i].Monat
      else:
        termine[i] = termine[i]._replace(Monat = '')
      if u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Reiseziel = u'Kuren an der polnischen Ostseeküste')
      if u'Kuren an der polnischen Ostseeküste' in termine[i-1].Reiseziel and u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel and termine[i-1].Tage == termine[i].Tage:
        termine[i] = termine[i]._replace(Termin = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reiseziel = termine[i].Reiseziel + ' (auch individuell buchbar)')
      for k in range(len(dates)):
        if termine[i].Veranstalter == 'RS':
          if 'Wander' in termine[i].Reiseziel:
            break
          if 'Kuren' in termine[i].Reiseziel:
            break
          if 'Musicalfahrt' in termine[i].Reiseziel:
            break
          if 'Elbphilharmonie' in termine[i].Reiseziel:
            break
          if termine[i].datum_beginn > dates[k] + timedelta(days=2):
            delta = termine[i].datum_beginn - dates[k]
            termine[i] = termine[i]._replace(tmp = delta.days-1)
            dates[k] = termine[i].datum_ende
            termine[i] = termine[i]._replace(bus = k+1)
            break
          if i > 0 and termine[i].datum_beginn <= dates[k] + timedelta(days=2) and k == (len(dates)-1):
            dates.append(termine[i].datum_ende)
            termine[i] = termine[i]._replace(bus = len(dates))
            termine[i] = termine[i]._replace(tmp = termine[i].tmp-1)
            break         

      #if (termine[i].anzahl_seiten_im_katalog == 0) and (termine[i].position_auf_seite == 1) and termine[i-1].Reiseziel != '':
        #termine[i-1] = termine[i-1]._replace(Reiseziel = termine[i-1].Reiseziel + ' (auch individuell buchbar)')
        #termine[i] = termine[i]._replace(Termin = '')

    cursor.close()

    ausgabeformat = request.GET.get('format')

    if ausgabeformat == 'html':
      return render(request, 'reisen/export_reiseterminuebersicht.html', {'termine': termine })

    elif ausgabeformat == 'timetable':

      return render(request, 'reisen/timetable.html', {'termine': termine })

    elif ausgabeformat == 'terminvorschau':

      return render(request, 'reisen/export_reiseterminuebersichtVORSCHAU.html', {'termine': termine })

    else:
      return render(request, 'reisen/export_reiseterminuebersicht.xml', {'termine': termine })

################################################################
# XML Export alle Reisen nach Zielregion/Kategorie Sommer 2017 #
################################################################
def reisezieluebersicht(request):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")

    #cursor.execute("select * from ( (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Kategorie' as kategorietyp, reisen_kategorie.kategorie, reisen_reise.korrektur_bemerkung_intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisekategorien on (reisen_reise.reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (reisen_reisekategorien.kategorie_id_id = reisen_kategorie.kategorieID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(reisen_reisepreiszusatz.preis, '.', ',')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2016 AND reisen_kategorie.kategorie in ('Wanderreisen', 'kombinierte Flug- und Busreisen', 'Flusskreuzfahrten') order by reisen_kategorie.kategorie, Seite) Union (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Zielregion' as kategorietyp, reisen_zielregion.name, reisen_reise.korrektur_bemerkung_intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisezielregionen on (reisen_reise.reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (reisen_reisezielregionen.zielregion_id_id = reisen_zielregion.zielregionID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(reisen_reisepreiszusatz.preis, '.', ',')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2016 AND reisen_zielregion.name in ('Deutschland', 'Benelux', 'Österreich / Schweiz', 'Tschechien / Slovakei / Ungarn / Rumänien', 'Frankreich', 'Italien', 'Slowenien / Kroatien / Montenegro / Bosnien und Herzegowina', 'Baltikum / Skandinavien / Finnland / Island', 'Polen', 'England / Schottland / Irland') order by reisen_zielregion.name, Seite)) as neuetabelle order by FIELD(neuetabelle.kategorie, 'Deutschland', 'Österreich / Schweiz', 'Benelux', 'Frankreich', 'Italien', 'Slowenien / Kroatien / Montenegro / Bosnien und Herzegowina', 'Tschechien / Slovakei / Ungarn / Rumänien', 'Polen', 'Baltikum / Skandinavien / Finnland / Island', 'England / Schottland / Irland', 'Wanderreisen', 'kombinierte Flug- und Busreisen', 'Flusskreuzfahrten'), katalogseite, position_auf_seite;")
    
    cursor.execute("select * from ( (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Kategorie' as kategorietyp, reisen_kategorie.kategorie, reisen_reise.korrektur_bemerkung_intern as intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisekategorien on (reisen_reise.reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (reisen_reisekategorien.kategorie_id_id = reisen_kategorie.kategorieID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2024 AND reisen_kategorie.kategorie in ('Wanderreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten') order by reisen_kategorie.kategorie, Seite) Union (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Zielregion' as kategorietyp, reisen_zielregion.name, reisen_reise.korrektur_bemerkung_intern as intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisezielregionen on (reisen_reise.reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (reisen_reisezielregionen.zielregion_id_id = reisen_zielregion.zielregionID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2024 AND reisen_zielregion.name in ('Deutschland', 'Österreich', 'Slowenien / Kroatien', 'Schweiz', 'Italien', 'Frankreich / Spanien / Andorra', 'Holland / Belgien', 'Britische Inseln', 'Skandinavien / Baltikum', 'Polen / Tschechien') order by reisen_zielregion.name, Seite)) as neuetabelle order by FIELD(neuetabelle.kategorie, 'Deutschland', 'Österreich', 'Slowenien / Kroatien', 'Schweiz', 'Italien', 'Frankreich / Spanien / Andorra', 'Holland / Belgien', 'Britische Inseln', 'Skandinavien / Baltikum', 'Polen / Tschechien', 'Wanderreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten'), katalogseite, position_auf_seite;")
    
    termine = namedtuplefetchall(cursor)

    kategorie = ''
    #swinemuende = False
    for i in range(len(termine)):
      if termine[i].kategorie != kategorie:
        kategorie = termine[i].kategorie
      else:
        termine[i] = termine[i]._replace(kategorie = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reise = termine[i].Reise + ' (auch individuell buchbar)')
      #if u'Swinemünde' in termine[i].Reise and swinemuende != True:
      #  termine[i] = termine[i]._replace(Tage = '8')
      #  swinemuende = True
      #elif u'Swinemünde' in termine[i].Reise and swinemuende != False:
      #  termine[i] = termine[i]._replace(Reise = '')
      #if (termine[i].anzahl_seiten_im_katalog == 0) and (termine[i].position_auf_seite == 1) and  termine[i-1].Reise is not None:
        #termine[i-1] = termine[i-1]._replace(Reise = termine[i-1].Reise + ' (auch individuell buchbar)')
        #termine[i] = termine[i]._replace(Reise = None)

    cursor.close()

    ausgabeformat = request.GET.get('format')
    if ausgabeformat == 'html':
      return render(request, 'reisen/export_reisezieluebersicht.html', {'termine': termine })
    else:
      return render(request, 'reisen/export_reisezieluebersicht.xml', {'termine': termine })

##################################################################
# Sommer 2017                                         #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer17(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = 'fa81408e2b69488498ace5b91737d187' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Benelux', u'Schweiz / Österreich / Tschechien / Slovakei / Ungarn', u'Frankreich / Italien / Andorra', u'Portugal', u'England / Schottland / Irland', u'Slowenien / Kroatien / Montenegro / Bosnien und Herzegowina', u'Baltikum / Skandinavien / Finnland / Island', u'Polen']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2018                                         #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer18(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '1ccf10f067e248e5ad079ce81991a64b' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Holland / Belgien', u'Österreich / Schweiz', u'Italien', u'Frankreich / Portugal / Spanien', u'Schottland / Irland', u'Slowenien / Kroatien / Montenegro / Bosnien und Herzegowina', u'Skandinavien / Baltikum / Finnland', u'Polen / Ukraine', u'Tschechien / Ungarn']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt18.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2019                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer19(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '45a66c434ad14ab1b65c298e5adb389a' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich / Schweiz', u'Holland / Belgien', u'Britische Inseln', u'Skandinavien / Baltikum / Finnland', u'Osteuropa', u'Balkan', u'Italien', u'Frankreich / Spanien / Portugal']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt19.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2020                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer20(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = 'c5c1b796502b4aa8aad0e73f01d161c7' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Rad- & Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich', u'Tschechien / Ungarn', u'Balkan', u'Italien / Schweiz', u'Frankreich / Spanien / Portugal / Andorra', u'Britische Inseln', u'Holland / Belgien', u'Skandinavien / Baltikum / Finnland', u'Polen / Russland']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt20.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2021                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer21(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = 'b393bbcb31204e6a9863b9e2ab91e708' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Rad- & Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich', u'Tschechien', u'Balkan', u'Italien', u'Frankreich / Spanien', u'Britische Inseln', u'Holland / Belgien', u'Nordeuropa', u'Polen / Russland']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt21.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2022                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer22(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '0ae00930970d470aae2ab6b871615801' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Rad- & Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich', u'Tschechien', u'Balkan', u'Italien', u'Frankreich / Spanien', u'Britische Inseln', u'Holland / Belgien', u'Nordeuropa', u'Polen / Russland']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt22.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2023                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer23(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '98968996adc544d483b28f8cb23444e7' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich / Ungarn', u'Tschechien', u'Slowenien / Kroatien', u'Schweiz', u'Italien', u'Frankreich', u'Britische Inseln', u'Holland / Belgien', u'Skandinavien', u'Polen']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt23.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2024                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer24(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = 'cb73c8d61acc49bf917dafbe6370bfb3' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich', u'Slowenien / Kroatien', u'Schweiz', u'Italien', u'Frankreich / Spanien / Andorra', u'Britische Inseln', u'Holland / Belgien', u'Skandinavien', u'Polen / Tschechien']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt24.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Sommer 2025                                                    #
##################################################################
#@python_2_unicode_compatible # For Python 3.4 and 2.7
def sommer25(request):
    dibug = ''
    cursor = connection.cursor()
    cursor.execute("SELECT DISTINCT reiseID, slug, sonstigeReisebeschreibung_titel, RP.abpreis, reisen_reise.neu, reisen_reise.individualbuchbar, reisen_kategorie.kategorie, reisen_zielregion.name, reisen_reise.untertitel, korrektur_bemerkung_intern, einleitung, reisetyp, KT.katalogseiten, RT.reise_id_id, reisen_reise.titel, RT.min_datum, RT.reisetermine FROM reisen_reise LEFT JOIN (SELECT reise_id_id, MIN(datum_beginn) AS min_datum, group_concat(DISTINCT CONCAT_WS(' - ', DATE_FORMAT(datum_beginn,'%d. %m. %Y'), DATE_FORMAT(datum_ende,'%d. %m. %Y')) ORDER BY datum_beginn ASC SEPARATOR '\n') AS reisetermine FROM reisen_reisetermine GROUP BY reise_id_id ORDER BY min_datum) AS RT ON (RT.reise_id_id = reiseID) LEFT JOIN (SELECT reise_id_id, group_concat(DISTINCT CONCAT_WS(' auf der Seite ', reisen_katalog.titel, katalogseite) ORDER BY position ASC SEPARATOR ' und im Katalog ') AS katalogseiten FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) GROUP BY reise_id_id) as KT ON (KT.reise_id_id = reiseID) LEFT JOIN reisen_reisekatalogzugehoerigkeit ON (reisen_reisekatalogzugehoerigkeit.reise_id_id = reisen_reise.reiseID) left join reisen_reisekategorien on (reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (kategorieID = kategorie_id_id) left join reisen_reisezielregionen on (reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (zielregionID = zielregion_id_id) left join (select reisen_reisepreise.reise_id_id, MIN(reisen_reisepreise.preis) AS abpreis from reisen_reisepreise GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reiseID = RP.reise_id_id) WHERE reisen_reisekatalogzugehoerigkeit.katalog_id_id = '972c50f33a7447be9f11fbfdceb85fc2' AND reisen_reise.status = 'f' ORDER BY RT.min_datum;")
    termine = namedtuplefetchall(cursor)
    cursor.close()

    kategorien = [ u'Wanderreisen', u'kombinierte Flug- & Busreisen', u'Kuren, Gesundheits- und Wellnessreisen', u'Flusskreuzfahrten']
    zielregionen = [ u'Deutschland', u'Österreich', u'Slowenien / Kroatien', u'Schweiz', u'Italien', u'Frankreich / Spanien / Andorra', u'Britische Inseln', u'Holland / Belgien', u'Skandinavien / Baltikum', u'Polen / Tschechien']

    dibug = ''#DefaultStorage().location + '  :::  ' + site.storage.location + site.directory + '  :::  ' + str(FileSystemStorage().directory_permissions_mode)

    termine_distinct = OrderedDict()
    for termin in termine:
      termine_distinct[termin.reiseID] = termin

    return render(request, 'reisen/index_sommerreisen_web_alt25.html', {'termine_distinct': termine_distinct.values(), 'termine': termine, 'dibug': dibug, 'kategorien': kategorien, 'zielregionen': zielregionen })

##################################################################
# Detail Seite, Reisedetails                                     #
##################################################################
def reise_detail(request, pk):

    #dibug = request.GET['dibug']
    dibug = ''

    #data = serializers.serialize("xml", Reisetermine.objects.filter(reise_id=pk).order_by('datum_beginn'))

    qs = Reisetermine.objects.filter(reise_id=pk).order_by('datum_beginn', 'datum_ende')

    #XMLSerializer = serializers.get_serializer("xml")
    #xml_serializer = XMLSerializer()
    #xml_serializer.serialize(qs)
    #data = xml_serializer.getvalue()

    #with open("./file.xml", "w") as out:
    #    xml_serializer.serialize(qs, stream=out)

    #dibug = 'NIX PASSIERT'
    reise = get_object_or_404(Reise, pk=pk)
    termine = Reisetermine.objects.filter(reise_id=pk).order_by('datum_beginn', 'datum_ende')
    abfahrtszeiten = Abfahrtszeiten.objects.filter(reise_id=pk).order_by('position')
    leistungen = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=0).order_by('position')
    leistungennichtindividual = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=1).order_by('position')
    zusatzleistungen = Zusatzleistung.objects.filter(reise_id=pk).order_by('position')
    fruehbucherrabatte = Fruehbucherrabatt.objects.filter(reise_id=pk).order_by('datum_bis')
    tage = Reisetage.objects.filter(reise_id=pk).order_by('tagnummer')
    reisebeschreibung = Reisebeschreibung.objects.filter(reise_id=pk).order_by('position')
    # Tagnummerntext erzeugen, bei Beschreibungen für mehrere Tage, Tagnummer x. - y. Tag erzeugen
    for idx, tag in enumerate(tage):
      tag.reisetagID = str(tag.reisetagID).replace('-','')
      if tag.tagnummertext:
        tag.nummerntext = tag.tagnummertext
      else:
        naechster_tag = tage[(idx+1) % len(tage)]
        if len(tage) == 1:
            tag.nummerntext = ''
        elif (naechster_tag.tagnummer == (tag.tagnummer + 1)) or (idx == (len(tage)-1)):
            tag.nummerntext = str(tag.tagnummer) + '. Tag:'
        else:
            tag.nummerntext = str(tag.tagnummer) + '. - ' + str(naechster_tag.tagnummer-1) + '. Tag:'
    #preise = Reisepreise.objects.filter(reise_id=pk).order_by('position')
    #preistitel = Preis.objects.filter(reise_id=pk).order_by('position')
    #preiszusatz = ReisepreisZusatz

    cursor = connection.cursor()
    cursor.execute("SELECT reise_id_id, reisepreisID, hauptpreis.titel, REPLACE(FORMAT(reisen_reisepreise.preis, 0),',','.') as preis, reisen_reisepreise.markierung, kommentar, subpreise.zpreis FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'), subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR '\n') as zpreis from reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisepreis_id_id) AS subpreise ON (reisepreisID = reisepreis_id_id) WHERE reisen_reisepreise.reise_id_id = '" + str(pk) + "' ORDER BY reisen_reisepreise.position;");
    preise = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT ausflugspaketID, ausflugspaket_text, reisetag_id_id, reisen_ausflugspakete.titel AS aptitel, erscheint_in, kommentar_titel, kommentar, reisen_ausflugspakete.position, reisen_preis.titel as ptitel, tagnummer, reisen_reisetage.titel as rtitel, preis, apleistungen.leistungen FROM reisen_ausflugspakete LEFT JOIN (SELECT ausflugspaket_id_id, group_concat(leistung ORDER BY position ASC SEPARATOR '\n') AS leistungen FROM reisen_leistungenausflugspaket group by ausflugspaket_id_id) AS apleistungen ON (reisen_ausflugspakete.ausflugspaketID = apleistungen.ausflugspaket_id_id) LEFT JOIN reisen_ausflugspaketpreise ON (reisen_ausflugspakete.ausflugspaketID = reisen_ausflugspaketpreise.ausflugspaket_id_id) LEFT JOIN reisen_preis ON (reisen_preis.preisID = reisen_ausflugspaketpreise.preis_id_id) LEFT JOIN reisen_ausflugspaketezureisetagen ON (reisen_ausflugspakete.ausflugspaketID = reisen_ausflugspaketezureisetagen.ausflugspaket_id_id) LEFT JOIN reisen_reisetage ON (reisen_reisetage.reisetagID = reisen_ausflugspaketezureisetagen.reisetag_id_id) WHERE reisen_ausflugspakete.reise_id_id = '" + str(pk) + "' ORDER BY reisen_ausflugspakete.position;");
    aps = namedtuplefetchall(cursor)
    cursor.close()

    aps_distinct = OrderedDict()
    zusatzleistungen_distinct = OrderedDict()
    for ap in aps:
        # wenn Leistung vorhanden dann ist es ein AP
        if ap.leistungen:
          aps_distinct[ap.ausflugspaketID] = ap
        # wenn nicht dann ist es eine Zusatzleistung
        else:
          zusatzleistungen_distinct[ap.ausflugspaketID] = ap

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisehinweise LEFT JOIN reisen_hinweis ON (reisen_reisehinweise.hinweis_id_id = reisen_hinweis.hinweisID) WHERE reisen_reisehinweise.reise_id_id = '" + str(pk) + "' ORDER BY reisen_reisehinweise.position;");
    hinweise = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisekategorien LEFT JOIN reisen_kategorie ON (reisen_reisekategorien.kategorie_id_id = reisen_kategorie.kategorieID) WHERE reisen_reisekategorien.reise_id_id = '" + str(pk) + "' ORDER BY reisen_reisekategorien.position;");
    kategorien = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisezielregionen LEFT JOIN reisen_zielregion ON (reisen_reisezielregionen.zielregion_id_id = reisen_zielregion.zielregionID) WHERE reisen_reisezielregionen.reise_id_id = '" + str(pk) + "' ORDER BY reisen_reisezielregionen.position;");
    zielregionen = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT bildID, zu_verwenden_in, bild, copyright, beschreibung, reisen_bild.titel as titel1, reisen_reisebilder.titel as titel2 FROM reisen_reisebilder LEFT JOIN reisen_bild ON (reisen_reisebilder.bild_id_id = reisen_bild.bildID) WHERE reisen_reisebilder.reise_id_id = '" + str(pk) + "' ORDER BY reisen_reisebilder.position;");
    bilder = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT reisen_katalog.katalogID, katalog_pdf, anzahl_seiten_im_katalog, katalogseite, reisen_katalog.titel FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) WHERE reisen_reisekatalogzugehoerigkeit.reise_id_id = '" + str(pk) + "' ORDER BY position;");
    kataloge = namedtuplefetchall(cursor)
    cursor.close()

    angebote = get_detail_angebote_queryset(pk)

    auftragsbestaetigungen = get_detail_auftragsbestaetigungen_queryset(pk)

    dibug = ''#request.GET.get('version')
    version = request.GET.get('version')
    kategorie_aktuell = request.GET.get('kategorie')
   
    #dibug = '' #querystring
    if version == 'alt':
        return render(
          request,
          'reisen/details_reisen_web_alt.html',
          {
              'reise': reise,
              'termine': termine,
              'abfahrtszeiten': abfahrtszeiten,
              'leistungen': leistungen,
              'leistungennichtindividual': leistungennichtindividual,
              'tage': tage,
              'reisebeschreibung': reisebeschreibung,
              'preise': preise,
              'aps': aps,
              'aps_distinct': aps_distinct.values(),
              'zusatzleistungen_distinct': zusatzleistungen_distinct.values(),
              'hinweise': hinweise,
              'kategorien': kategorien,
              'kategorie_aktuell': kategorie_aktuell,
              'zielregionen': zielregionen,
              'zusatzleistungen': zusatzleistungen,
              'fruehbucherrabatte': fruehbucherrabatte,
              'bilder': bilder,
              #'bild_str': bild_str,
              'kataloge': kataloge,
              'angebote': angebote,
              'auftragsbestaetigungen': auftragsbestaetigungen,
              'dibug': dibug,
          }
      )
    elif version == 'docx':

      orte = {
        'HBF': u'Hbf. Schwerin',
        'VSB': u'Hast. v. Stauffenberg Str.',
        'GAR': u'Gartenstadt',
        'WIS': u'ZOB Wismar',
        'ROG': 'Gadebusch Roggendorfer Str.',
        'ANK': u'Ankunft zurück in Schwerin'
      }

      locale.setlocale(locale.LC_TIME, "de_DE")

      document = Document()
      core_properties = document.core_properties
      core_properties.author = 'Andreas Finger'
      core_properties.language = 'Deutsch'
      core_properties.category = 'Reisebeschreibung'
      core_properties.comments = 'Reisebeschreibung Sommerreisen 2018'
      core_properties.content_status = 'Draft'
      core_properties.created = reise.datum_erzeugung
      core_properties.identifier = reise.reiseID
      core_properties.keywords = "Sommerreisen"
      core_properties.last_modified_by = reise.zuletzt_bearbeitet_von_id
#      core_properties.title = re.sub(r'\xb7|\xdf|\xf6|\u2019', '-', reise.titel)
      core_properties.title = re.sub(r'[^a-zA-Z0-9_ ."-]', '-', reise.titel)
      
      docx_title="reise_" + str(reise.reiseID) + ".docx"
      #document.add_picture('/var/www/reiseservice-schwerin/rss2' + static('images/logo_google_blau.png'), width=Inches(2))
      #sections = document.sections
      #sections[0].orientation = WD_ORIENT.LANDSCAPE
      #document.sections[0].orientation = WD_ORIENT.LANDSCAPE
      font = document.styles['Normal'].font
      font.name = 'Submariner R24'
      font.size = Pt(12)
      paragraph_format = document.styles['Normal'].paragraph_format
      paragraph_format.space_before = Pt(0)
      paragraph_format.space_after = Pt(0)
      paragraph_format.line_spacing = 1

      if reise.korrektur_bemerkung_intern:
        p = document.add_paragraph()
        run = p.add_run()
        f = run.font
        f.color.rgb = RGBColor(0xff, 0x00, 0x00)
        run.add_text(reise.korrektur_bemerkung_intern)
        p.add_run("\n")

      for katalog in kataloge:
        p = document.add_paragraph()
        run = p.add_run()
        run.italic = True
        run.add_text("Diese Reise finden Sie im Druckexemplar unseres Kataloges " + katalog.titel + " auf Seite " + str(katalog.katalogseite) + ".")
        p.add_run("\n")

      p = document.add_paragraph()
      p.add_run(reise.reisetyp).bold = True
      if reise.individualreisetext or reise.individualreisetitel:
        if reise.individualbuchbar:
          p = document.add_paragraph()
          run = p.add_run()
          run.italic = True
          run.add_text(reise.individualbuchbar)
        else:
          p = document.add_paragraph()
          run = p.add_run()
          run.italic = True
          run.add_text("Auch mit individuellem Aufenthalt buchbar.")
      
      if reise.veranstalter == 'RS':
        p = document.add_paragraph()
      if reise.veranstalter == 'SH':
        p = document.add_paragraph()
        p.add_run('Veranstalter: Sewert Reisen').italic = True
        p.add_run("\n")
      elif reise.veranstalter == 'RT':
        p = document.add_paragraph()
        p.add_run(unicode('Veranstalter: R&T Reisen Ludwigslust',"utf-8")).italic = True
        p.add_run("\n")

      if reise.zubucher:
        p = document.add_paragraph()
        p.add_run(reise.zubucher).italic = True
        p.add_run("\n")

      p = document.add_paragraph()
      run = p.add_run()
      run.bold = True
      run.font.size = Pt(16)
      run.add_text(reise.titel)
      
      p = document.add_paragraph()
      run = p.add_run()
      run.bold = True
      run.add_text(reise.untertitel)
      
      if reise.einleitung:
        p = document.add_paragraph()
        p.add_run("\n")
        p.add_run(reise.einleitung)

      p = document.add_paragraph()
      p.add_run("\n")
      if len(termine)>1:
        p.add_run('Reisetermine:').bold = True
      else:
        p.add_run('Reisetermin:').bold = True

      for termin in termine:
        if termin.kommentar:
          p = document.add_paragraph()
          p.add_run(termin.kommentar).italic = True
        p = document.add_paragraph()
        if termin.datum_ende:
          p.add_run(datetime.strftime(termin.datum_beginn, "%d.%m.%y") + ' - ' + datetime.strftime(termin.datum_ende, "%d.%m.%y") + " " + termin.markierung)
        else:
          p.add_run(datetime.strftime(termin.datum_beginn, "%d.%m.%y") + " " + termin.markierung)

      if abfahrtszeiten:
        p = document.add_paragraph()
        p.add_run("\n")
        if len(abfahrtszeiten)>1:
          p.add_run('Abfahrtszeiten:').bold = True
        else:
          p.add_run('Abfahrtszeit:').bold = True
        for abfahrtszeit in abfahrtszeiten:
          if abfahrtszeit.zeit:
            p = document.add_paragraph()
            p.add_run(str(abfahrtszeit.zeit)[:5] + ' Uhr ' + orte[abfahrtszeit.ort] + abfahrtszeit.kommentar)
          else:
            p = document.add_paragraph()
            p.add_run("Die genaue Zeit für " + orte[abfahrtszeit.ort] + " erfragen Sie bitte im Reisebüro.")
      elif reise.reisetyp == 'Tagesfahrt':
            p = document.add_paragraph()
            p.add_run("Die Abfahrtszeiten sind noch nicht bekannt!")
      
      if preise:
        p = document.add_paragraph()
        p.add_run("\n")
        if len(preise)>1:
          p.add_run('Preise:').bold = True
        else:
          p.add_run('Preis:').bold = True
        for preis in preise:
          if preis.kommentar == 'ab':
            p = document.add_paragraph()
            p.add_run(preis.titel + ": " + preis.kommentar + " " + str(preis.preis) + unicode(" € ", "utf-8"))
          else:
            p = document.add_paragraph()
            p.add_run(preis.markierung + preis.titel + ": " + preis.preis + unicode(" € ", "utf-8"))
            if preis.kommentar:
              p.add_run(preis.kommentar)
          if preis.zpreis:
            p = document.add_paragraph()
            p.add_run(preis.zpreis)
          p.add_run("\n")
      else:
        p = document.add_paragraph()
        p.add_run("Preis:")
        p.add_run("\n")
        p.add_run("Der Preis ist noch nicht bekannt.")
        p.add_run("\n")

      if fruehbucherrabatte:
        p = document.add_paragraph()
        if len(fruehbucherrabatte)>1:
          p.add_run(unicode("Fruehbucherrabatte:","utf-8")).bold = True
        else:
          p.add_run(unicode("Fruehbucherrabatt:","utf-8")).bold = True
        for fruehbucherrabatt in fruehbucherrabatte:
          p = document.add_paragraph()
          p.add_run(str(fruehbucherrabatt.rabatt) + unicode(" € ", "utf-8") + fruehbucherrabatt.rabattbezeichnung + " bis zum " + datetime.strftime(fruehbucherrabatt.datum_bis, "%d.%m.%y"))
        p.add_run("\n")

      if reisebeschreibung:
        p = document.add_paragraph()
        p.add_run("allgemeine Beschreibungen:").bold = True
        for beschreibung in reisebeschreibung:
          p = document.add_paragraph()
          p.add_run(beschreibung.titel).bold = True
          p = document.add_paragraph()
          p.add_run(beschreibung.beschreibung + " " + beschreibung.zusatz)
        p.add_run("\n")

      if tage:
        p = document.add_paragraph()
        p.add_run("Reiseablauf:").bold = True
        for tag in tage:
          if aps:
            p = document.add_paragraph()
            p.add_run(tag.nummerntext + " " + tag.titel).bold = True
            for ap in aps:
              if ap.reisetag_id_id == tag.reisetagID:
                if ap.erscheint_in == 'Titel':
                  p.add_run(ap.ausflugspaket_text)
            for ap in aps:
              if ap.reisetag_id_id == tag.reisetagID:
                if ap.erscheint_in == 'Text':
                  p = document.add_paragraph()
                  p.add_run(ap.ausflugspaket_text)
            p = document.add_paragraph()
            run = p.add_run()
            run.add_text(tag.beschreibung + " ")
            run = p.add_run()
            run.bold = True
            run.add_text(tag.zusatz)
          else:
            p = document.add_paragraph()
            p.add_run(tag.nummerntext + " " + tag.titel).bold = True
            p = document.add_paragraph()
            run = p.add_run()
            run.add_text(tag.beschreibung + " ")
            run = p.add_run()
            run.bold = True
            run.add_text(tag.zusatz)
          p = document.add_paragraph()

      if hinweise:
        p = document.add_paragraph()
        if len(hinweise)>1:
          p.add_run('Hinweise:').bold = True
        else:
          p.add_run('Hinweis:').bold = True
        for hinweis in hinweise:
          p = document.add_paragraph()
          p.add_run(hinweis.hinweis).italic = True

      if reise.individualreisetitel or reise.individualreisetext:
        p = document.add_paragraph()
        p.add_run("\n")
        if reise.individualreisetitel:
          p.add_run(reise.individualreisetitel + ": ").bold = True
        if reise.individualreisetext:
          p.add_run(reise.individualreisetext)

      if leistungen or leistungennichtindividual:
        p = document.add_paragraph()
        p.add_run("\n")
        if len(leistungen)>1 or len(leistungennichtindividual)>1:
          p.add_run('Leistungen:').bold = True
        else:
          p.add_run('Leistung:').bold = True
        for leistung in leistungen:
          p = document.add_paragraph()
          p.add_run("- " + leistung.leistung)
        for leistungnichtindividual in leistungennichtindividual:
          p = document.add_paragraph()
          p.add_run("* " + leistungnichtindividual.leistung)
        if reise.leistungen_kommentar:
          p = document.add_paragraph()
          p.add_run("\n")
          p.add_run(reise.leistungen_kommentar).italic = True

      if aps_distinct:
        p = document.add_paragraph()
        p.add_run("\n")
        if len(aps_distinct)>1:
          p.add_run('Ausflugspakete:').bold = True
        else:
          p.add_run('Ausflugspaket:').bold = True
        for ap in aps_distinct.values():
          if ap.aptitel:
            p = document.add_paragraph()
            p.add_run(ap.aptitel).bold = True
          if ap.preis:
            p = document.add_paragraph()
            p.add_run("Preis: " + str(ap.preis) + unicode(" €", "utf-8"))
          else:
            p = document.add_paragraph()
            p.add_run("Der Preis ist noch nicht bekannt.")
          if ap.kommentar_titel:
            p = document.add_paragraph()
            p.add_run(ap.kommentar_titel).italic = True
          p = document.add_paragraph()
          p.add_run("enthaltene Leistungen:")
          p = document.add_paragraph()
          p.add_run(ap.leistungen)
          if ap.kommentar:
            p = document.add_paragraph()
            p.add_run(ap.kommentar).italic = True

      if zusatzleistungen_distinct:
        p = document.add_paragraph()
        p.add_run("\n")
        if reise.zusatzleistungen_titel:
          p.add_run(reise.zusatzleistungen_titel)
        else:
          if len(zusatzleistungen_distinct)>1:
            p.add_run(unicode("zusätzlich angebotene Leistungen:","utf-8")).bold = True
          else:
            p.add_run(unicode("zusätzlich angebotene Leistung:","utf-8")).bold = True
        if reise.zusatzleistungen_kommentar:
          p = document.add_paragraph()
          p.add_run(reise.zusatzleistungen_kommentar)
        for zusatzleistung in zusatzleistungen_distinct.values():
          p = document.add_paragraph()
          p.add_run("- " + zusatzleistung.aptitel)
          if zusatzleistung.preis:
            p.add_run("\t" + str(zusatzleistung.preis) + unicode(" €","utf-8"))
          else:
            p.add_run("\t Der Preis ist noch nicht bekannt.")
          if zusatzleistung.kommentar_titel:
            p = document.add_paragraph()
            p.add_run(zusatzleistung.kommentar_titel)
          if zusatzleistung.kommentar:
            p = document.add_paragraph()
            p.add_run(zusatzleistung.kommentar)
          if reise.zusatzleistungen_fuss_kommentar:
            p = document.add_paragraph()
            p.add_run(reise.zusatzleistungen_fuss_kommentar)


#      for tagesfahrt in tagesfahrten:
#        p = document.add_paragraph()
#        p.add_run(datetime.strftime(tagesfahrt['datum_beginn'], "%a., %d.%m.") + "\t" + tagesfahrt['reisedaten'].titel).bold = True
#        if tagesfahrt['reisedaten'].untertitel or tagesfahrt['reisedaten'].einleitung:
#          p = document.add_paragraph()
#        if tagesfahrt['reisedaten'].untertitel:
#          p.add_run(tagesfahrt['reisedaten'].untertitel + ' ')
#        if  tagesfahrt['reisedaten'].einleitung:
#          p.add_run(tagesfahrt['reisedaten'].einleitung)
#        if tagesfahrt['reisedaten'].veranstalter != 'RS':
#          if tagesfahrt['reisedaten'].veranstalter == 'SH':
#            p = document.add_paragraph()
#            p.add_run('Veranstalter: Sewert Reisen').italic = True
#        if tagesfahrt['hinweise']:
#          for hinweis in tagesfahrt['hinweise']:
#            p = document.add_paragraph()
#            p.add_run(unicode(hinweis['hinweis'], "utf-8")).italic = True
#        for idx, abfahrtszeit in enumerate(tagesfahrt['abfahrtszeiten']):
#          if idx == 0:
#            p = document.add_paragraph()
#            p.add_run('Abfahrt: ').name = 'Submariner R24 light'
#          if idx == (len(tagesfahrt['abfahrtszeiten'])-1):
#            p = document.add_paragraph()
#            p.add_run('Ankunft: ').name = 'Submariner R24 light'
#          p.add_run(str(abfahrtszeit.zeit)[:5] + ' Uhr ' + orte[abfahrtszeit.ort]).name = 'Submariner R24 light'
#          if idx < (len(tagesfahrt['abfahrtszeiten'])-2):
#            p.add_run(', ').name = 'Submariner R24 light'
#
#        newparagraph = False
#        for idx, preis in enumerate(tagesfahrt['preise']):
#          if preis['markierung']:
#            if preis['markierung'] == tagesfahrt['datum_markierung']:
#              if newparagraph:
#                p = document.add_paragraph()
#              newparagraph = True
#              p.add_run('\t\t' + unicode(preis['preistitel'], "utf-8") + ': ').bold = True
#              if preis['kommentar']:
#                p.add_run(preis['kommentar'] + ' ').bold = True
#              p.add_run(str(preis['preis']) + u' €').bold = True
#          else:
#            if newparagraph:
#              p = document.add_paragraph()
#            newparagraph = True
#            p.add_run('\t\t' + unicode(preis['preistitel'], "utf-8") + ': ').bold = True
#            if preis['kommentar']:
#              p.add_run(preis['kommentar'] + ' ').bold = True
#            p.add_run(str(preis['preis']) + u' €').bold = True
#
#        p = document.add_paragraph()
#        p.add_run().size = Pt(8)

      # -----------------------------
      # Prepare document for download
      # -----------------------------
      f = StringIO()
      document.save(f)
      length = f.tell()
      f.seek(0)
      response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
      )
      response['Content-Disposition'] = 'attachment; filename=' + docx_title
      response['Content-Length'] = length
      return response

    else:
        return render(
          request,
          'reisen/reise_detail.html',
          {
              'reise': reise,
              'termine': termine,
              'abfahrtszeiten': abfahrtszeiten,
              'leistungen': leistungen,
              'leistungennichtindividual': leistungennichtindividual,
              'tage': tage,
              'reisebeschreibung': reisebeschreibung,
              'preise': preise,
              'aps': aps,
              'aps_distinct': aps_distinct.values(),
              'zusatzleistungen_distinct': zusatzleistungen_distinct.values(),
              'hinweise': hinweise,
              'kategorien': kategorien,
              'zielregionen': zielregionen,
              'zusatzleistungen': zusatzleistungen,
              'fruehbucherrabatte': fruehbucherrabatte,
              'bilder': bilder,
              'kataloge': kataloge,
              'angebote': angebote,
              'auftragsbestaetigungen': auftragsbestaetigungen,
              'dibug': dibug,
          }
      )

##################################################################
# Detail Seite, Reisedetails                                     #
##################################################################
def reise_detail_export(request, pk):

    dibug = ''
    reise = get_object_or_404(Reise, pk=pk)
    termine = Reisetermine.objects.filter(reise_id=pk).order_by('datum_beginn', 'datum_ende')
    abfahrtszeiten = Abfahrtszeiten.objects.filter(reise_id=pk).order_by('position')

    for abfahrtszeit in abfahrtszeiten:
      if abfahrtszeit.ort == 'HBF':
        abfahrtszeit.ort = 'Hbf. Schwerin'
      elif abfahrtszeit.ort == 'VSB':
        abfahrtszeit.ort = 'HST "v. Stauffenberg Str."'
      elif abfahrtszeit.ort == 'ANK':
        abfahrtszeit.ort = 'Ankunft'
      elif abfahrtszeit.ort == 'GAR':
        abfahrtszeit.ort = 'Gartenstadt'
      elif abfahrtszeit.ort == 'ROG':
        abfahrtszeit.ort = 'Gadebusch Roggendorfer Str.'

    leistungen = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=0, leistungkurhotel=0).order_by('position')
    leistungennichtindividual = LeistungenReise.objects.filter(reise_id=pk).filter(nichtindividual=1).order_by('position')
    leistungenkurhotel = LeistungenReise.objects.filter(reise_id=pk).filter(leistungkurhotel=1).order_by('position')
    zusatzleistungen = Zusatzleistung.objects.filter(reise_id=pk).order_by('position')
    fruehbucherrabatte = Fruehbucherrabatt.objects.filter(reise_id=pk).order_by('datum_bis')
    tage = Reisetage.objects.filter(reise_id=pk).order_by('tagnummer')
    reisebeschreibung = Reisebeschreibung.objects.filter(reise_id=pk).order_by('position')
    # Tagnummerntext erzeugen, bei Beschreibungen für mehrere Tage, Tagnummer x. - y. Tag erzeugen
    pkquery = re.sub(r'[-]', '', str(pk))

    for idx, tag in enumerate(tage):
      	tag.reisetagID = str(tag.reisetagID).replace('-','')
        if tag.tagnummertext:
          tag.nummerntext = tag.tagnummertext + ' '
        else:
          naechster_tag = tage[(idx+1) % len(tage)]
          if len(tage) == 1:
              tag.nummerntext = ''
          elif (naechster_tag.tagnummer == (tag.tagnummer + 1)) or (idx == (len(tage)-1)):
              tag.nummerntext = str(tag.tagnummer) + '. Tag: '
          else:
              tag.nummerntext = str(tag.tagnummer) + '. - ' + str(naechster_tag.tagnummer-1) + '. Tag: '

    cursor = connection.cursor()
    cursor.execute("SELECT reise_id_id, reisepreisID, hauptpreis.titel, REPLACE(FORMAT(reisen_reisepreise.preis, 0),',','.') as preis, reisen_reisepreise.markierung, kommentar, subpreise.zpreis FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(':	', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'), subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR '\n') as zpreis from reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisepreis_id_id) AS subpreise ON (reisepreisID = reisepreis_id_id) WHERE reisen_reisepreise.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisepreise.position;");
    preise = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT ausflugspaketID, ausflugspaket_text, reisetag_id_id, reisen_ausflugspakete.titel AS aptitel, erscheint_in, kommentar_titel, kommentar, reisen_ausflugspakete.position, reisen_preis.titel as ptitel, tagnummer, reisen_reisetage.titel as rtitel, preis, apleistungen.leistungen FROM reisen_ausflugspakete LEFT JOIN (SELECT ausflugspaket_id_id, group_concat(leistung ORDER BY position ASC SEPARATOR '\n') AS leistungen FROM reisen_leistungenausflugspaket group by ausflugspaket_id_id) AS apleistungen ON (reisen_ausflugspakete.ausflugspaketID = apleistungen.ausflugspaket_id_id) LEFT JOIN reisen_ausflugspaketpreise ON (reisen_ausflugspakete.ausflugspaketID = reisen_ausflugspaketpreise.ausflugspaket_id_id) LEFT JOIN reisen_preis ON (reisen_preis.preisID = reisen_ausflugspaketpreise.preis_id_id) LEFT JOIN reisen_ausflugspaketezureisetagen ON (reisen_ausflugspakete.ausflugspaketID = reisen_ausflugspaketezureisetagen.ausflugspaket_id_id) LEFT JOIN reisen_reisetage ON (reisen_reisetage.reisetagID = reisen_ausflugspaketezureisetagen.reisetag_id_id) WHERE reisen_ausflugspakete.reise_id_id = '" + pkquery + "' ORDER BY reisen_ausflugspakete.position;");
    aps = namedtuplefetchall(cursor)
    cursor.close()

    aps_distinct = OrderedDict()
    zusatzleistungen_distinct = OrderedDict()
    for ap in aps:
        # wenn Leistung vorhanden dann ist es ein AP
        if ap.leistungen:
          aps_distinct[ap.ausflugspaketID] = ap
        # wenn nicht dann ist es eine Zusatzleistung
        else:
          zusatzleistungen_distinct[ap.ausflugspaketID] = ap

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisehinweise LEFT JOIN reisen_hinweis ON (reisen_reisehinweise.hinweis_id_id = reisen_hinweis.hinweisID) WHERE reisen_reisehinweise.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisehinweise.position;");
    hinweise = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisekategorien LEFT JOIN reisen_kategorie ON (reisen_reisekategorien.kategorie_id_id = reisen_kategorie.kategorieID) WHERE reisen_reisekategorien.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisekategorien.position;");
    kategorien = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    cursor.execute("SELECT * FROM reisen_reisezielregionen LEFT JOIN reisen_zielregion ON (reisen_reisezielregionen.zielregion_id_id = reisen_zielregion.zielregionID) WHERE reisen_reisezielregionen.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisezielregionen.position;");
    zielregionen = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    #cursor.execute("SELECT bild, beschreibung, reisen_bild.titel as titel1, reisen_reisebilder.titel as titel2 FROM reisen_reisebilder LEFT JOIN reisen_bild ON (reisen_reisebilder.bild_id_id = reisen_bild.bildID) WHERE reisen_reisebilder.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisebilder.position;");
    cursor.execute("SELECT REPLACE(bild,'images/','') as bild, beschreibung, reisen_bild.titel as titel1, reisen_reisebilder.titel as titel2, reisen_bildanbieter.bildanbieter as bildanbieter, bildnummer, copyright, url as bildurl, kommentar as bildkommentar FROM reisen_reisebilder LEFT JOIN reisen_bild ON (reisen_reisebilder.bild_id_id = reisen_bild.bildID) LEFT JOIN reisen_bildanbieter ON (reisen_bild.bildanbieter_id_id = reisen_bildanbieter.bildanbieterID) WHERE reisen_reisebilder.reise_id_id = '" + pkquery + "' ORDER BY reisen_reisebilder.position;");
    bilder = namedtuplefetchall(cursor)
    cursor.close()

    cursor = connection.cursor()
    #cursor.execute("SELECT katalog_pdf, anzahl_seiten_im_katalog, katalogseite, reisen_katalog.titel FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) WHERE reisen_reisekatalogzugehoerigkeit.reise_id_id = '" + str(pk) + "' ORDER BY position;");
    cursor.execute("SELECT anzahl_seiten_im_katalog, katalogseite, reisen_katalog.titel, position_auf_seite FROM reisen_reisekatalogzugehoerigkeit LEFT JOIN reisen_katalog ON (reisen_katalog.katalogID = reisen_reisekatalogzugehoerigkeit.katalog_id_id) WHERE reisen_reisekatalogzugehoerigkeit.reise_id_id = '" + pkquery + "'"); # AND reisen_katalog.katalogID = '1f9e043d4f4341a8a563d408fd267c56'");
    katalog = namedtuplefetchall(cursor)
    cursor.close()

    bilder_anzeigen = request.GET.get('bilder')

    #dibug = '' #querystring
    if bilder_anzeigen == 'ja':
        return render(
          request,
          'reisen/reise_detail_export_mit_bildern.xml',
          {
          'reise': reise,
          'termine': termine,
          'abfahrtszeiten': abfahrtszeiten,
          'leistungen': leistungen,
          'leistungennichtindividual': leistungennichtindividual,
          'tage': tage,
          'reisebeschreibung': reisebeschreibung,
          'preise': preise,
          'aps': aps,
          'aps_distinct': aps_distinct.values(),
          'zusatzleistungen_distinct': zusatzleistungen_distinct.values(),
          'hinweise': hinweise,
          'kategorien': kategorien,
          'zielregionen': zielregionen,
          'zusatzleistungen': zusatzleistungen,
          'fruehbucherrabatte': fruehbucherrabatte,
          'bilder': bilder,
          'katalog': katalog[0],
          'dibug': dibug,
          }
        )
    elif bilder_anzeigen == 'nur':
      return render(
        request,
        'reisen/reise_detail_export_nur_bilder.xml',
        {
          'reise': reise,
          'termine': termine,
          'abfahrtszeiten': abfahrtszeiten,
          'leistungen': leistungen,
          'leistungennichtindividual': leistungennichtindividual,
          'tage': tage,
          'reisebeschreibung': reisebeschreibung,
          'preise': preise,
          'aps': aps,
          'aps_distinct': aps_distinct.values(),
          'zusatzleistungen_distinct': zusatzleistungen_distinct.values(),
          'hinweise': hinweise,
          'kategorien': kategorien,
          'zielregionen': zielregionen,
          'zusatzleistungen': zusatzleistungen,
          'fruehbucherrabatte': fruehbucherrabatte,
          'bilder': bilder,
          'katalog': katalog[0],
          'dibug': dibug,
        }
      )
    else:
      return render(
        request,
        'reisen/reise_detail_export.xml',
        {
          'reise': reise,
          'termine': termine,
          'abfahrtszeiten': abfahrtszeiten,
          'leistungen': leistungen,
          'leistungennichtindividual': leistungennichtindividual,
          'leistungenkurhotel': leistungenkurhotel,
          'tage': tage,
          'reisebeschreibung': reisebeschreibung,
          'preise': preise,
          'aps': aps,
          'aps_distinct': aps_distinct.values(),
          'zusatzleistungen_distinct': zusatzleistungen_distinct.values(),
          'hinweise': hinweise,
          'kategorien': kategorien,
          'zielregionen': zielregionen,
          'zusatzleistungen': zusatzleistungen,
          'fruehbucherrabatte': fruehbucherrabatte,
          'bilder': bilder,
          'katalog': katalog[0],
          'dibug': dibug,
        }
      )

#####################################################################
# XML Export alle Reisen nach Zielregion/Kategorie für Gesamtexport #
#####################################################################
def reisezieluebersicht_alles(pk):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")

    cursor.execute("select * from ( (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Kategorie' as kategorietyp, reisen_kategorie.kategorie, reisen_reise.korrektur_bemerkung_intern as intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisekategorien on (reisen_reise.reiseID = reisen_reisekategorien.reise_id_id) left join reisen_kategorie on (reisen_reisekategorien.kategorie_id_id = reisen_kategorie.kategorieID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '" + str(pk) + "' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2020 AND reisen_kategorie.kategorie in ('Wanderreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten') order by reisen_kategorie.kategorie, Seite) Union (select DISTINCT reisen_reise.reiseID, reisen_reise.individualbuchbar, reisen_reise.neu, 'Zielregion' as kategorietyp, reisen_zielregion.name, reisen_reise.korrektur_bemerkung_intern as intern, reisen_reise.titel as Reise, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite, RP.preise from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) left join reisen_reisezielregionen on (reisen_reise.reiseID = reisen_reisezielregionen.reise_id_id) left join reisen_zielregion on (reisen_reisezielregionen.zielregion_id_id = reisen_zielregion.zielregionID) left join (SELECT reisen_reisepreise.reise_id_id, GROUP_CONCAT(CONCAT(hauptpreis.titel, ': ', IF(reisen_reisepreise.kommentar = 'ab', 'ab ', ''), concat(replace(FORMAT(reisen_reisepreise.preis, 0),',','.'), ' €'), reisen_reisepreise.markierung, IF(reisen_reisepreise.kommentar = 'ab', ', ', IF(reisen_reisepreise.kommentar = '', ', ', concat(', ', reisen_reisepreise.kommentar, ', '))), IFNULL(subpreise.zpreis,'')) ORDER BY reisen_reisepreise.position ASC SEPARATOR ', ') as preise FROM reisen_reisepreise LEFT JOIN reisen_preis AS hauptpreis ON (hauptpreis.preisID = reisen_reisepreise.preis_id_id) LEFT JOIN (SELECT reisen_reisepreiszusatz.reisepreis_id_id, GROUP_CONCAT(IF(reisen_reisepreiszusatz.preis > 0, CONCAT(CONCAT_WS(': ', subpreis.titel, replace(FORMAT(reisen_reisepreiszusatz.preis, 0),',','.')), ' €'),  subpreis.titel) ORDER BY reisen_reisepreiszusatz.position ASC SEPARATOR ', ') as zpreis FROM reisen_reisepreiszusatz LEFT JOIN reisen_preis AS subpreis ON (subpreis.preisID = reisen_reisepreiszusatz.preis_id_id) GROUP BY reisen_reisepreiszusatz.reisepreis_id_id) AS subpreise ON (reisen_reisepreise.reisepreisID = subpreise.reisepreis_id_id) GROUP BY reisen_reisepreise.reise_id_id) AS RP ON (reisen_reise.reiseID = RP.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '" + str(pk) + "' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2020 AND reisen_zielregion.name in ('Deutschland', 'Österreich', 'Slowenien / Kroatien', 'Schweiz', 'Italien', 'Frankreich / Spanien / Andorra', 'Holland / Belgien', 'Britische Inseln', 'Skandinavien / Baltikum', 'Polen / Tschechien') order by reisen_zielregion.name, Seite)) as neuetabelle order by FIELD(neuetabelle.kategorie, 'Deutschland', 'Österreich', 'Slowenien / Kroatien', 'Schweiz', 'Italien', 'Frankreich / Spanien / Andorra', 'Holland / Belgien', 'Britische Inseln', 'Skandinavien / Baltikum', 'Polen / Tschechien', 'Wanderreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten'), katalogseite, position_auf_seite;")

    termine = namedtuplefetchall(cursor)

    kategorie = ''
    swinemuende = False
    for i in range(len(termine)):
      if termine[i].kategorie != kategorie:
        kategorie = termine[i].kategorie
      else:
        termine[i] = termine[i]._replace(kategorie = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reise = termine[i].Reise + ' (auch individuell buchbar)')
      if u'Swinemünde' in termine[i].Reise and swinemuende != True:
        termine[i] = termine[i]._replace(Tage = '5/8')
        swinemuende = True
      elif u'Swinemünde' in termine[i].Reise and swinemuende != False:
        termine[i] = termine[i]._replace(Reise = '')
      #if (termine[i].anzahl_seiten_im_katalog == 0) and (termine[i].position_auf_seite == 1) and  termine[i-1].Reise is not None:
        #termine[i-1] = termine[i-1]._replace(Reise = termine[i-1].Reise + ' (auch individuell buchbar)')
        #termine[i] = termine[i]._replace(Reise = None)

    cursor.close()

    return termine

########################################################
# XML Export alle Reisen nach Termin fuer Gesamtexport #
########################################################
def reiseterminuebersicht_alles(pk):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")
    cursor.execute("select reisen_reise.reiseID as pk, reisen_reise.korrektur_bemerkung_intern as intern, Veranstalter, DAYOFYEAR(datum_beginn) as tmp, 0 as tmp2, 0 as bus, date_format(reisen_reisetermine.datum_beginn, '%M') as Monat, reisen_reise.individualbuchbar, reisen_reise.neu, reisen_reise.titel as Reiseziel, reisen_reisetermine.datum_beginn, reisen_reisetermine.datum_ende, CONCAT(date_format(reisen_reisetermine.datum_beginn,'%d.%m.'), '-', date_format(reisen_reisetermine.datum_ende,'%d.%m.%y')) as Termin, (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage, katalogseite, anzahl_seiten_im_katalog, position_auf_seite, concat(katalogseite,if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')) as Seite from reisen_reise left join reisen_reisetermine on (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) left join reisen_reisekatalogzugehoerigkeit on(reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) where reisen_reisekatalogzugehoerigkeit.katalog_id_id = '" + str(pk) + "' AND date_format(reisen_reisetermine.datum_beginn, '%Y') > 2019 order by datum_beginn, datum_ende, katalogseite, position_auf_seite;")

    termine = namedtuplefetchall(cursor)

    month = ''
    dates = []
    delta = 0
    for i in range(len(termine)):
      if i == 0:
        dates.append(termine[i].datum_ende)
        termine[i] = termine[i]._replace(tmp = termine[i].tmp-1)
        if termine[i].Veranstalter == 'RS':
          termine[i] = termine[i]._replace(bus = 1)
      if termine[i].Monat != month:
        month = termine[i].Monat
      else:
        termine[i] = termine[i]._replace(Monat = '')
      if u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Reiseziel = u'Kuren an der polnischen Ostseeküste')
      if u'Kuren an der polnischen Ostseeküste' in termine[i-1].Reiseziel and u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel and termine[i-1].Tage == termine[i].Tage:
        termine[i] = termine[i]._replace(Termin = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reiseziel = termine[i].Reiseziel + ' (auch individuell buchbar)')
      for k in range(len(dates)):
        if termine[i].Veranstalter == 'RS':
          if 'Wander' in termine[i].Reiseziel:
            break
          if 'Musicalfahrt' in termine[i].Reiseziel:
            break
          if 'Elbphilharmonie' in termine[i].Reiseziel:
            break          
          if termine[i].datum_beginn > dates[k] + timedelta(days=2):
            delta = termine[i].datum_beginn - dates[k]
            termine[i] = termine[i]._replace(tmp = delta.days-1)
            dates[k] = termine[i].datum_ende
            termine[i] = termine[i]._replace(bus = k+1)
            break
          if i > 0 and termine[i].datum_beginn <= dates[k] + timedelta(days=2) and k == (len(dates)-1):
            dates.append(termine[i].datum_ende)
            termine[i] = termine[i]._replace(bus = len(dates))
            termine[i] = termine[i]._replace(tmp = termine[i].tmp-1)
            break

    cursor.close()

    return termine

######################################################################
# XML Export alle Reisen Winterkatalog nach Termin fuer Gesamtexport #
######################################################################
def reiseuebersichtwinter_alles(pk):

    cursor = connection.cursor()
    cursor.execute("SET lc_time_names = 'de_DE';")
    cursor.execute("SELECT   reiseID, reisen_reise.korrektur_bemerkung_intern as intern, kat.groupkat as k,   reisen_reise.individualbuchbar,   reisen_reise.neu,   reisen_reise.titel as Reiseziel,   reisen_reise.sonstigeReisebeschreibung_titel as Zusatztitel,   CONCAT(     date_format(reisen_reisetermine.datum_beginn,'%d.%m.'),     '-',     date_format(reisen_reisetermine.datum_ende,'%d.%m.%y')   ) as Termin,   (to_days(reisen_reisetermine.datum_ende)-to_days(reisen_reisetermine.datum_beginn)+1) as Tage,   concat(     katalogseite,     if(anzahl_seiten_im_katalog>1,concat('|',katalogseite+1),'')   ) as Seite FROM   reisen_reise LEFT JOIN   reisen_reisetermine   ON   (reisen_reise.reiseID = reisen_reisetermine.reise_id_id) LEFT JOIN (SELECT reise_id_id, group_concat(kategorie) as groupkat from reisen_reisekategorien left join reisen_kategorie on kategorieID = kategorie_id_id where kategorie in ('kombinierte Flug- & Busreisen', 'Flusskreuzfahrten', 'Busreisen') group by reise_id_id) AS kat ON (kat.reise_id_id = reisen_reise.reiseID) LEFT JOIN   reisen_reisekatalogzugehoerigkeit   ON   (reisen_reise.reiseID = reisen_reisekatalogzugehoerigkeit.reise_id_id) WHERE   reisen_reisekatalogzugehoerigkeit.katalog_id_id = '" + str(pk) + "' AND katalogseite > 0 ORDER BY FIELD(k, 'Busreisen', 'kombinierte Flug- & Busreisen', 'Flusskreuzfahrten'),   datum_beginn,   datum_ende,   katalogseite,   position_auf_seite;")
    termine = namedtuplefetchall(cursor)

    swine = False
    skifahren = False
    gruenewoche = False
    barcelona = False
    danke = False
    blau = False
    azoren = False
    rom = False
    andalusien = False
    lissabon = False
    starlight = False
    elfi = False
    wellnesswochenende = False
    kategorie = ''
    for i in range(len(termine)):
      if termine[i].k != kategorie:
        kategorie = termine[i].k
      else:
        termine[i] = termine[i]._replace(k = '')
      if u'Weihnachten und Silvester an der polnischen Ostseeküste' in termine[i-1].Reiseziel and u'Weihnachten und Silvester an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Kuren an der polnischen Ostseeküste' in termine[i].Reiseziel:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Blauf' in termine[i].Reiseziel and blau != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25')
        blau = True
      elif u'Blauf' in termine[i].Reiseziel and blau != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Starlight' in termine[i].Reiseziel and starlight != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & 26')
        starlight = True
      elif u'Starlight' in termine[i].Reiseziel and starlight != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Elbphil' in termine[i].Reiseziel and elfi != True:
        termine[i] = termine[i]._replace(Termin = 'Dez. 25 & Apr. 26')
        elfi = True
      elif u'Elbphil' in termine[i].Reiseziel and elfi != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Wellnesswochenende' in termine[i].Reiseziel and wellnesswochenende != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        wellnesswochenende = True
      elif u'Wellnesswochenende' in termine[i].Reiseziel and wellnesswochenende != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Dankesch' in termine[i].Reiseziel and danke != True:
        termine[i] = termine[i]._replace(Termin = 'Nov. 25 & Jan. 26')
        danke = True
      elif u'Dankesch' in termine[i].Reiseziel and danke != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Azoren' in termine[i].Reiseziel and azoren != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 22 & Apr. 23')
        azoren = True
      elif u'Azoren' in termine[i].Reiseziel and azoren != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'ewige Stadt' in termine[i].Reiseziel and rom != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & 26')
        rom = True
      elif u'ewige Stadt' in termine[i].Reiseziel and rom != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Andalusien' in termine[i].Reiseziel and andalusien != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        andalusien = True
      elif u'Andalusien' in termine[i].Reiseziel and andalusien != False:
        termine[i] = termine[i]._replace(Termin = '')
      #if u'Swinemünde' in termine[i].Reiseziel and swine != True:
      #  termine[i] = termine[i]._replace(Termin = 'Okt. 25 bis März 26')
      #  termine[i] = termine[i]._replace(Tage = '8')
      #  swine = True
      #elif u'Swinemünde' in termine[i].Reiseziel and swine != False:
      #  termine[i] = termine[i]._replace(Termin = '')
      if u'Barcelona' in termine[i].Reiseziel and barcelona != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & März 26')
        barcelona = True
      elif u'Barcelona' in termine[i].Reiseziel and barcelona != False:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Lissabon' in termine[i].Reiseziel and lissabon != True:
        termine[i] = termine[i]._replace(Termin = 'Okt. 25 & Mai 26')
        lissabon = True
      elif u'Lissabon' in termine[i].Reiseziel and lissabon != False:
        termine[i] = termine[i]._replace(Termin = '')
      #if u'Skifahren in Südtirol' in termine[i].Reiseziel and skifahren != True:
      #  termine[i] = termine[i]._replace(Termin = 'Jan. & Feb. 23')
      #  skifahren = True
      #elif u'Skifahren in Südtirol' in termine[i].Reiseziel and skifahren != False:
      #  termine[i] = termine[i]._replace(Termin = '')
      if not termine[i].Tage:
        termine[i] = termine[i]._replace(Termin = '')
      if u'Grünen Woche' in termine[i].Reiseziel and gruenewoche == False:
        termine[i] = termine[i]._replace(Termin = '20., 22., 24.01.26')
        termine[i] = termine[i]._replace(Tage = '1')
        gruenewoche = True
      elif u'Grünen Woche' in termine[i].Reiseziel and gruenewoche == True:
        termine[i] = termine[i]._replace(Termin = '')
      if termine[i].individualbuchbar != '':
        termine[i] = termine[i]._replace(Reiseziel = termine[i].Reiseziel + ' (auch individuell buchbar)')
      #if (termine[i].anzahl_seiten_im_katalog == 0) and (termine[i].position_auf_seite == 1) and termine[i-1].Reiseziel != '':
        #termine[i-1] = termine[i-1]._replace(Reiseziel = termine[i-1].Reiseziel + ' (auch individuell buchbar)')
        #termine[i] = termine[i]._replace(Termin = '')

    cursor.close()

    return termine

##################################################################
# Export aller Reisen eines Kataloges                            #
##################################################################
def reise_detail_export_alles(request, pk):

    dibug = ''#Reisekatalogzugehoerigkeit.objects.filter(katalog_id=pk,reise_id='2522a98b79de42699534a0b0f306a95c').values('katalogseite')
    reiseids = Reise.objects.select_related().filter(
        reisekatalogzugehoerigkeit__katalog_id=pk,
#        status='f',
      ).order_by('reisekatalogzugehoerigkeit__katalogseite').distinct()

    reisen = [
      {
        "reiseID": rid.reiseID,
        "titel": rid.titel,
        "untertitel": rid.untertitel,
        "einleitung": rid.einleitung,
        "sonstigeReisebeschreibung_titel": rid.sonstigeReisebeschreibung_titel,
        "neu": rid.neu,
        "reisetyp": rid.reisetyp,
        "zubucher": rid.zubucher,
        "veranstalter": rid.veranstalter,
        "individualbuchbar": rid.individualbuchbar,
        "individualreisetitel": rid.individualreisetitel,
        "individualreisetext": rid.individualreisetext,
        "leistungen_kopfkommentar": rid.leistungen_kopfkommentar,
        "leistungen_kommentar": rid.leistungen_kommentar,
        "zusatzleistungen_titel": rid.zusatzleistungen_titel,
        "zusatzleistungen_kommentar": rid.zusatzleistungen_kommentar,
        "zusatzleistungen_fuss_kommentar": rid.zusatzleistungen_fuss_kommentar,
        "katalogseite": Reisekatalogzugehoerigkeit.objects.filter(katalog_id=pk,reise_id=rid.reiseID).values_list('katalogseite', flat=True)[0],
        "anzahl_seiten_im_katalog": Reisekatalogzugehoerigkeit.objects.filter(katalog_id=pk,reise_id=rid.reiseID).values_list('anzahl_seiten_im_katalog', flat=True)[0],
        "position_auf_seite": Reisekatalogzugehoerigkeit.objects.filter(katalog_id=pk,reise_id=rid.reiseID).values_list('position_auf_seite', flat=True)[0],
        "leistungen": LeistungenReise.objects.filter(reise_id=rid.reiseID).filter(nichtindividual=0,leistungkurhotel=0).order_by('position'),
        "leistungennichtindividual": LeistungenReise.objects.filter(reise_id=rid.reiseID).filter(nichtindividual=1).order_by('position'),
        "leistungenkurhotel": LeistungenReise.objects.filter(reise_id=rid.reiseID).filter(leistungkurhotel=1).order_by('position'),
        "reisebeschreibung": Reisebeschreibung.objects.filter(reise_id=rid.reiseID).order_by('position'),
        "abfahrtszeiten": Abfahrtszeiten.objects.filter(reise_id=rid.reiseID).order_by('position'),
        "tage": Reisetage.objects.filter(reise_id=rid.reiseID).order_by('tagnummer'),
        #"katalogseite": list(Reisekatalogzugehoerigkeit.objects.filter(katalog_id=pk).values('katalogseite')),
        #"reisedaten": get_object_or_404(Reise, pk=rid.reiseID),
        "termine": Reisetermine.objects.filter(reise_id=rid.reiseID).order_by('datum_beginn', 'datum_ende'),
        "kategorien": [
          {
            "kategorie": str(kategorie.kategorie_id)
          }
          for kategorie in Reisekategorien.objects.filter(reise_id=rid.reiseID)
        ],
        "zielregionen": [
          {
            "name": str(zielregion.zielregion_id)
          }
          for zielregion in Reisezielregionen.objects.filter(reise_id=rid.reiseID)
        ],
        "hinweise": [
          {
            "hinweis": str(hinweis.hinweis_id)
          }
          for hinweis in Reisehinweise.objects.filter(reise_id=rid.reiseID)
        ],
        "ausflugspakete": [
          {
            "titel": ausflugspaket.titel,
            "kommentar_titel": ausflugspaket.kommentar_titel,
            "kommentar": ausflugspaket.kommentar,
            "preis": (Ausflugspaketpreise.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position').values_list('preis', flat=True)[0] if len(Ausflugspaketpreise.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position').values_list('preis', flat=True)) > 0 else None),
            "preis_titel": (Preis.objects.select_related().filter(ausflugspaketpreise__ausflugspaket_id=ausflugspaket.ausflugspaketID).values_list('titel', flat=True)[0] if len(Preis.objects.select_related().filter(ausflugspaketpreise__ausflugspaket_id=ausflugspaket.ausflugspaketID).values_list('titel', flat=True)) > 0 else None),
            "leistungen": [
              {
                "leistung": leistung.leistung
              }
              for leistung in LeistungenAusflugspaket.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position')
            ]
          }
          for ausflugspaket in Ausflugspakete.objects.filter(reise_id=rid.reiseID).order_by('position') if LeistungenAusflugspaket.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID)
        ],
        "zusatzleistungen": [
          {
            "titel": ausflugspaket.titel,
            "kommentar_titel": ausflugspaket.kommentar_titel,
            "kommentar": ausflugspaket.kommentar,
            "preis": (Ausflugspaketpreise.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position').values_list('preis', flat=True)[0] if len(Ausflugspaketpreise.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position').values_list('preis', flat=True)) > 0 else None),
            "preis_titel": (Preis.objects.select_related().filter(ausflugspaketpreise__ausflugspaket_id=ausflugspaket.ausflugspaketID).values_list('titel', flat=True)[0] if len(Preis.objects.select_related().filter(ausflugspaketpreise__ausflugspaket_id=ausflugspaket.ausflugspaketID).values_list('titel', flat=True)) > 0 else None),
            "leistungen": [
              {
                "leistung": leistung.leistung
              }
              for leistung in LeistungenAusflugspaket.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID).order_by('position')
            ]
          }
          for ausflugspaket in Ausflugspakete.objects.filter(reise_id=rid.reiseID).order_by('position') if not LeistungenAusflugspaket.objects.filter(ausflugspaket_id=ausflugspaket.ausflugspaketID)
        ],
        "preise": [
          {
            "preis": preis.preis,
            "kommentar": preis.kommentar,
            "markierung": preis.markierung,
            "titel": str(preis.preis_id),
            "zusatzpreise": [
              {
                "zusatzpreis_titel": str(zusatzpreis.preis_id),
                "zusatzpreis": zusatzpreis.preis #wenn preis 0 .. weitere preise auf anfrage ... als fußkommentar zu zusatzpreisen?? ham we sowas?
              }
              for zusatzpreis in ReisepreisZusatz.objects.filter(reisepreis_id=preis.reisepreisID)
            ]
          }
          for preis in Reisepreise.objects.filter(reise_id=rid.reiseID).order_by('position')
        ],
        "fruehbucherrabatte": [
          {
            "rabattbezeichnung" : fruehbucherrabatt.rabattbezeichnung,
            "datum_bis": fruehbucherrabatt.datum_bis,
            "rabatt": fruehbucherrabatt.rabatt
          }
          for fruehbucherrabatt in Fruehbucherrabatt.objects.filter(reise_id=rid.reiseID).order_by('datum_bis')
        ]
      }
      for rid in reiseids
    ]

    for reise in reisen:
      # Tagnummerntext erzeugen, bei Beschreibungen für mehrere Tage, Tagnummer x. - y. Tag erzeugen
      for idx, tag in enumerate(reise['tage']):
        tag.reisetagID = str(tag.reisetagID).replace('-','')
        if tag.tagnummertext:
          tag.nummerntext = tag.tagnummertext + ' '
        else:
          naechster_tag = reise['tage'][(idx+1) % len(reise['tage'])]
          if len(reise['tage']) == 1:
              tag.nummerntext = ''
          elif (naechster_tag.tagnummer == (tag.tagnummer + 1)) or (idx == (len(reise['tage'])-1)):
              tag.nummerntext = str(tag.tagnummer) + '. Tag: '
          else:
              tag.nummerntext = str(tag.tagnummer) + '. - ' + str(naechster_tag.tagnummer-1) + '. Tag: '
      # Haltestellen korrigieren
      for abfahrtszeit in reise['abfahrtszeiten']:
        if abfahrtszeit.ort == 'HBF':
          abfahrtszeit.ort = 'HBF'
        elif abfahrtszeit.ort == 'VSB':
          abfahrtszeit.ort = 'v. Stauffenberg Str.'
        elif abfahrtszeit.ort == 'ANK':
          abfahrtszeit.ort = 'Ankunft'
        elif abfahrtszeit.ort == 'GAR':
          abfahrtszeit.ort = 'Gartenstadt'
        elif abfahrtszeit.ort == 'ROG':
          abfahrtszeit.ort = 'Gadebusch Roggendorfer Str.'

    katalog = request.GET.get('katalog')
    if katalog == 'winter':
      template = 'reisen/reise_detail_export_alles_winter.xml'
      return render(
        request,
        template,
        {
          'dibug': dibug,
          'reisen': list(reisen),
          'termine_reiseuebersichtwinter': reiseuebersichtwinter_alles(pk),
        }
      )  
    else:
      template = 'reisen/reise_detail_export_alles.xml'

      return render(
        request,
        template,
        {
          'dibug': dibug,
          'reisen': list(reisen),
          'termine_reisezieluebersicht': reisezieluebersicht_alles(pk),
          'termine_reiseterminuebersicht': reiseterminuebersicht_alles(pk),
        }
      )

##################################################################
# Zustiege                                                       #
##################################################################
def zustiege(request):

    conn = mysql.connector.connect(
      host="localhost",
      user="django",
      database="deltaplan",
      passwd="MMu9U30iL!"
    )

    #conn = MySQLdb.connect("connection info here")
    try:
      cursor = conn.cursor()
      cursor.execute("select Ort, Preis from Zustiege where Ort like '%Haustür%' order by Ort")
      zustiege = cursor.fetchall()
    finally:
      conn.close()

    return render(request, 'reisen/zustiege.html', {'zustiege': zustiege })

##################################################################
# aktuelle Reisen                                                #
##################################################################
def aktuelleReisen(request):

    conn = mysql.connector.connect(
      host="localhost",
      user="django",
      database="deltaplan",
      passwd="MMu9U30iL!"
    )

    #conn = MySQLdb.connect("connection info here")
    try:
      cursor = conn.cursor()
      cursor.execute("select Ziel, date_format(vom,'%d.%m.%Y') as start, date_format(bis,'%d.%m.%Y') as ende, Preis from Reisen where vom > curdate() and storno is null and vom != bis order by vom limit 10")
      aktuelleReisen = cursor.fetchall()
    finally:
      conn.close()

    return render(request, 'reisen/aktuelleReisen.html', {'aktuelleReisen': aktuelleReisen })

##################################################################
# Zubringer                                                      #
##################################################################
def zubringer(request):

    conn = mysql.connector.connect(
      host="localhost",
      user="django",
      database="deltaplan",
      passwd="MMu9U30iL!"
    )

    rnr = request.GET.get('rnr')
    jahr = request.GET.get('jahr')

    #conn = MySQLdb.connect("connection info here")
    try:
      cursor = conn.cursor()

      anfrage = """
        SELECT
          Namen.Anrede AS Anrede,
          Namen.Vorname AS Vorname,
          Namen.Name AS Name,
          IF(OptionBis LIKE '1899-12-30 00:00:00', '', 'Option') AS Reservierung,
          IF(Warteliste = 0, '', 'Warteliste') AS Warteliste,
          Namen.Bnr AS Bnr,
          Namen.ZOrt AS Ortkurz,
          Zustiege.Ort AS Ort,
          IF(
            CHAR_LENGTH(Namen.ZZeit) < 4 AND CHAR_LENGTH(Namen.ZZeit) > 1,
            CONCAT('0',LEFT(ZZEIT,1),':',RIGHT(Namen.ZZeit,2),' Uhr'),
            IF(
              CHAR_LENGTH(Namen.ZZeit) < 2,
              'keine Zeit vergeben',
              CONCAT(LEFT(Namen.ZZeit,2), ':', RIGHT(Namen.ZZeit,2), ' Uhr')
            )
          ) AS Zeit
          FROM
            Namen
          LEFT JOIN
            Buchungen
          ON
            Namen.Jahr = Buchungen.Jahr
          AND
            Namen.Bnr = Buchungen.Bnr
          LEFT JOIN
            Zustiege
          ON
            Namen.ZOrt = Zustiege.Kurz
          WHERE
            Buchungen.Storno IS NULL
          AND
            Zustiege.Preis = 0
          AND
            Namen.Jahr = """ + jahr + """
          AND
            Namen.Rnr = """ + rnr + """
          ORDER BY
            Namen.ZZeit,
            Namen.ZOrt,
            Namen.Bnr,
            Namen.PaxNr
          """

      cursor.execute(anfrage)
      uebergabe = cursor.fetchall()

      uebergabeliste = [
        {
          "id": i,
          "anrede": ub[0],
          "vorname": ub[1],
          "name": ub[2],
          "option": ub[3],
          "warteliste": ub[4],
          "selbebnr": True if i > 0 and ub[5] == uebergabe[i-1][5] else False,
          "bnr": ub[5],
          "selberort": True if i > 0 and ub[6] == uebergabe[i-1][6] else False,
          "ortkurz": ub[6],
          "ort": ub[7],
          "zeit": ub[8],
        }
        for i, ub in enumerate(uebergabe)
      ]

      anfrage = """
        SELECT
          Namen.Anrede AS Anrede,
          Namen.Vorname AS Vorname,
          Namen.Name AS Name,
          IF(OptionBis like '1899-12-30 00:00:00', '', 'Option') AS Reservierung,
          IF(Warteliste = 0, '', 'Warteliste') AS Warteliste,
          Namen.Bnr AS Bnr,
          Namen.ZOrt AS Ortkurz,
          Zustiege.Ort AS Ort,
          IF(
            CHAR_LENGTH(Namen.ZZeit) < 4 AND CHAR_LENGTH(Namen.ZZeit) > 1,
            CONCAT('0',LEFT(ZZEIT,1),':',RIGHT(Namen.ZZeit,2),' Uhr'),
            IF(
              CHAR_LENGTH(Namen.ZZeit) < 2,
              'keine Zeit vergeben',
              CONCAT(LEFT(Namen.ZZeit,2), ':', RIGHT(Namen.ZZeit,2), ' Uhr')
            )
          ) AS Zeit,
          Taxis.Name1,
          Taxis.Name2,
          Taxis.Strasse,
          Taxis.Ort,
          Taxis.Telefon1,
          Taxis.Telefon2,
          Taxis.Telefax,
          Taxis.email
          FROM
            Namen
          LEFT JOIN
            Buchungen
          ON
            Namen.Jahr = Buchungen.Jahr
          AND
            Namen.Bnr = Buchungen.Bnr
          LEFT JOIN
            Zustiege
          ON
            Namen.ZOrt = Zustiege.Kurz
          LEFT JOIN
            Taxis
          ON
            Zustiege.Taxi = Taxis.Kurz
          WHERE
            Buchungen.Storno IS NULL
          AND
            (Zustiege.Preis > 0 OR Zustiege.Preis IS NULL)
          AND
            Namen.Jahr = """ + jahr + """
          AND
            Namen.Rnr = """ + rnr + """
          ORDER BY
            Namen.ZZeit,
            Namen.ZOrt,
            Namen.Bnr,
            Namen.PaxNr
          """

      cursor.execute(anfrage)
      zubringer = cursor.fetchall()

      zubringerliste = [
        {
          "id": i,
          "anrede": zb[0],
          "vorname": zb[1],
          "name": zb[2],
          "option": zb[3],
          "warteliste": zb[4],
          "selbebnr": True if i > 0 and zb[5] == zubringer[i-1][5] else False,
          "bnr": zb[5],
          "selberort": True if i > 0 and zb[6] == zubringer[i-1][6] else False,
          "ortkurz": zb[6],
          "ort": zb[7],
          "zeit": zb[8],
          "taxiname1": zb[9],
          "taxiname2": zb[10],
          "taxistrasse": zb[11],
          "taxiort": zb[12],
          "taxitel1": zb[13],
          "taxitel2": zb[14],
          "taxifax": zb[15],
          "taxiemail": zb[16],
        }
        for i, zb in enumerate(zubringer)
      ]

      anfrage = """
        SELECT
          Ziel,
          DATE_FORMAT(vom, '%d.%m.%Y') AS vom,
          DATE_FORMAT(bis, '%d.%m.%Y') AS bis,
          rnr,
          jahr,
          Nurzustiege
        FROM
          Reisen
        WHERE
          Jahr = """ + jahr + """
        AND
          Rnr = """ + rnr

      cursor.execute(anfrage)
      reise = cursor.fetchall()

      anfrage = "select Kurz, Taxi, Ort from Zustiege where Preis = 0"
     
      cursor.execute(anfrage)
      zustiege = cursor.fetchall()

      reiseinfos = [
        {
          "id": i,
          "ziel": ri[0],
          "vom": ri[1],
          "bis": ri[2],
          "rnr": ri[3],
          "jahr": ri[4],
          "zustiege": [
            {
              "zustiegkurz": zustieg.split(u"\xb0")[1] if len(zustieg.split(u"\xb0")) > 1 else None,
              "zustieg": None, #u"PLACEHOLDER" if len(zustieg.split(u"\xb0")) > 1 else None,
              "zeit": zustieg.split(u"\xb0")[0] if len(zustieg.split(u"\xb0")) > 1 else None
            }
            for k, zustieg in enumerate(ri[5][1:].split(u"\n"))
          ]
        }
        for i, ri in enumerate(reise)
      ]

      for m, ri in enumerate(reiseinfos):
        for j, z in enumerate(ri['zustiege']):
          zkurz0 = z['zustiegkurz']
          for n, zustieg in enumerate(zustiege):
            zkurz1 = zustieg[0]
            if zkurz0 is not None and zkurz0 == zkurz1:
              z['zustieg'] = zustieg[2]
   
    finally:
      conn.close()

    return render(request, 'reisen/zubringer.html', {'jahr': jahr, 'rnr': rnr, 'uebergabeliste': uebergabeliste, 'zubringerliste': zubringerliste, 'reiseinfos': reiseinfos })


##################################################################
# Zustiege                                                       #
##################################################################
def neu(request):

    conn = mysql.connector.connect(
      host="localhost",
      user="django",
      database="deltaplan",
      passwd="MMu9U30iL!"
    )

    #conn = MySQLdb.connect("connection info here")
    try:
      cursor = conn.cursor()
      cursor.execute("select Ort, Preis from Zustiege where Ort like '%Haustür%' order by Ort")
      zustiege = cursor.fetchall()
    finally:
      conn.close()

    return render(request, 'reisen/neu.html', {'neu': neu })

#@python_2_unicode_compatible # For Python 3.4 and 2.7
@csrf_exempt
def printZubringer(request):

    if request.method == "POST":
      reihenfolge = request.POST['reihenfolge']
      uebergabeliste = request.POST['uebergabeliste']
      zubringerliste = request.POST['zubringerliste']
      jahr = request.POST['jahr']
      rnr = request.POST['rnr']
      reiseziel = HTMLParser().unescape(request.POST['reiseziel'])
      vom = request.POST['vom']
      bis = request.POST['bis']
      lis = request.POST['lis']

    try:
      rf = json.loads('[{"' + reihenfolge.replace('[]=','":').replace('&','},{"') + '}]')
      ul = json.loads(uebergabeliste.replace('&#39;', '"').replace(': u"', ': "').replace('None','"None"').replace('False','"False"').replace('True','"True"'))
      zl = json.loads(zubringerliste.replace('&#39;', '"').replace(': u"', ': "').replace('None','"None"').replace('False','"False"').replace('True','"True"'))
    except (ValueError):
      rf = "No JSON object could be decoded"
      ul = "No JSON object could be decoded"
      zl = "No JSON object could be decoded"

    debug = "TEST"

    locale.setlocale(locale.LC_TIME, "de_DE")

    document = Document()
    core_properties = document.core_properties
    core_properties.author = 'Andreas Finger'
    core_properties.language = 'Deutsch'
    core_properties.category = 'Zustiegsliste'
    core_properties.comments = 'Zustiege und Zubringer in richtiger Reihenfolge'
    core_properties.content_status = 'Draft'
    #core_properties.created = date.today()
    #core_properties.created = ("%s" % datetime.strftime(datetime.now(), "%B"))
    core_properties.identifier = jahr + '/' + rnr
    core_properties.keywords = "Zubringer"
    core_properties.last_modified_by = "af"
#    core_properties.title = re.sub(r'\xb7|\xdf|\xf6|\u2019', '-', reise.titel)
    core_properties.title = re.sub(r'[^a-zA-Z0-9_ ."-]', '-', reiseziel)
#    core_properties.title = reiseziel

    #docx_title="reise_" + str(reiseziel) + ".docx"
    docx_title="reise_" + core_properties.title + ".docx"
    #document.add_picture('/var/www/reiseservice-schwerin/rss2' + static('images/logo_google_blau.png'), width=Inches(2))
    #sections = document.sections
    #sections[0].orientation = WD_ORIENT.LANDSCAPE
    #document.sections[0].orientation = WD_ORIENT.LANDSCAPE
    font = document.styles['Normal'].font
    font.name = 'Submariner R24'
    font.size = Pt(12)
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1

    p = document.add_paragraph()
    run = p.add_run()
    f = run.font
    f.color.rgb = RGBColor(0xff, 0x00, 0x00)
    f.size = Pt(16)
    run.add_text(u'Zubringerliste R' + rnr + '/' + jahr + ': ' + reiseziel + ' vom ' + vom + ' bis ' + bis)
    p.add_run("\n")
    #run.add_text(uebergabeliste)
    #p.add_run("\n")
    #run.add_text(str(ul))
    #p.add_run("\n")

#    for katalog in kataloge:
    p = document.add_paragraph()
    #run = p.add_run()
    #run.italic = True
    #run.add_text(u"Reihenfolge: ")
    #p.add_run("\n")
    for i, index in enumerate(rf):
      k = index.values()[0]
      try:
        if index.keys()[0] == "uebergabe":
          p.add_run("\n")
          run = p.add_run()
          f = run.font
          f.size = Pt(16)
          run.bold = True
          #run.add_text(index.keys()[0].replace("uebergabe","Bushalt").replace("zubringer","Zubringer"))
          #run.add_text(" ")
          #run.add_text(str(index.values()[0]))
          #run.add_text(": ")
          run.add_text(ul[k-1]['ort'] + ' um ' + ul[k-1]['zeit'])
          p.add_run("\n")
          while k < len(ul) and eval(ul[k]['selberort']):
            #p.add_run("\n")
            #p.add_run(str(k)  + ' - ' + str(len(ul)))
            if not eval(ul[k-1]['selbebnr']):
              p.add_run("\n")
              p.add_run('B' + str(ul[k-1]['bnr']))
              if (ul[k-1]['option']):
                p.add_run(' (' + ul[k-1]['option'] + ')')
              if (ul[k-1]['warteliste']):
                p.add_run(' (' + ul[k-1]['warteliste'] + ')')
              p.add_run("\n")
            #p.add_run(str(k-1)+ ' - ' + str(ul[k-1]['id']) + ': ' + ul[k-1]['name'] + ', ' + ul[k-1]['vorname'])
            p.add_run(ul[k-1]['anrede'] + ' ' + ul[k-1]['name'] + ', ' + ul[k-1]['vorname'])
            p.add_run("\n")
            k += 1
          #p.add_run("\n")
          #p.add_run(str(k)  + ' - ' + str(len(ul)))
          if not eval(ul[k-1]['selbebnr']):
            p.add_run("\n")
            p.add_run('B' + str(ul[k-1]['bnr']))
            p.add_run("\n")
          #p.add_run(str(k-1)+ ' - ' + str(ul[k-1]['id']) + ': ' + ul[k-1]['name'] + ', ' + ul[k-1]['vorname'])
          p.add_run(ul[k-1]['anrede'] + ' ' + ul[k-1]['name'] + ', ' + ul[k-1]['vorname'])
          p.add_run("\n")
        else:
          p.add_run("\n")
          run = p.add_run()
          run.bold = True
          run.add_text("Taxi aus ")
          run.add_text(zl[k-1]['ort'] + ' mit Abfahrt um ' + zl[k-1]['zeit'])
          p.add_run("\n")
          while k < len(zl) and eval(zl[k]['selberort']):
            #p.add_run("\n")
            #p.add_run(str(k)  + ' - ' + str(len(zl)))
            if not eval(zl[k-1]['selbebnr']):
              p.add_run("\n")
              p.add_run('B' + str(zl[k-1]['bnr']))
              p.add_run("\n")
            #p.add_run(str(k-1)+ ' - ' + str(zl[k-1]['id']) + ': ' + zl[k-1]['name'] + ', ' + zl[k-1]['vorname'])
            p.add_run(zl[k-1]['anrede'] + ' ' + zl[k-1]['name'] + ', ' + zl[k-1]['vorname'])
            p.add_run("\n")
            k += 1
          #p.add_run("\n")
          #p.add_run(str(k)  + ' - ' + str(len(zl)))
          if not eval(zl[k-1]['selbebnr']):
            p.add_run("\n")
            p.add_run('B' + str(zl[k-1]['bnr']))
            p.add_run("\n")
          #p.add_run(str(k-1)+ ' - ' + str(zl[k-1]['id']) + ': ' + zl[k-1]['name'] + ', ' + zl[k-1]['vorname'])
          p.add_run(zl[k-1]['anrede'] + ' ' + zl[k-1]['name'] + ', ' + zl[k-1]['vorname'])
          p.add_run("\n")
      except (IndexError):
        p.add_run("INDEXERROR: " + str(k))

    # -----------------------------
    # Prepare document for download
    # -----------------------------
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
      f.getvalue(),
      content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    response["Content-Encoding"] = "UTF-8"

    data = {
       'successmsg': 'SUCCESS',
       'debug': debug,
       'reihenfolge': reihenfolge,
       'uebergabeliste': uebergabeliste,
       'worddoc': document,
       'error': 'ERROR'
     }

    return response
    #return JsonResponse(data)

##################################################################
# Ablauf                                                         #
##################################################################
def ablauf(request):

    conn = mysql.connector.connect(
      host="localhost",
      user="django",
      database="deltaplan",
      passwd="MMu9U30iL!"
    )

    rnr = request.GET.get('rnr')
    jahr = request.GET.get('jahr')

    docx_title = rnr + u" / " + jahr

    #conn = MySQLdb.connect("connection info here")
    try:
      cursor = conn.cursor()

      anfrage = """
        SELECT
          Reisen.Ziel AS Ziel,
          DATE_FORMAT(Reisen.vom, '%d.%m.%Y') AS vom,
          DATE_FORMAT(Reisen.bis, '%d.%m.%Y') AS bis,
          DATEDIFF(Reisen.bis, Reisen.vom) AS Dauer
        FROM
          Reisen
        WHERE
          Jahr = """ + jahr + """
        AND
          rnr = """ + rnr

      cursor.execute(anfrage)
      titel = cursor.fetchall()

      for t in titel:
        ziel = t[0]
        vom = t[1]
        bis = t[2]
        dauer = t[3]

      anfrage = """
        SELECT
          Namen.Anrede AS Anrede,
          Namen.Vorname AS Vorname,
          Namen.Name AS Name,
          IF(OptionBis LIKE '1899-12-30 00:00:00', '', 'Option') AS Reservierung,
          IF(Warteliste = 0, '', 'Warteliste') AS Warteliste,
          Namen.Bnr AS Bnr,
          Namen.ZOrt AS Ortkurz,
          Zustiege.Ort AS Ort,
          IF(
            CHAR_LENGTH(Namen.ZZeit) < 4 AND CHAR_LENGTH(Namen.ZZeit) > 1,
            CONCAT('0',LEFT(ZZEIT,1),':',RIGHT(Namen.ZZeit,2),' Uhr'),
            IF(
              CHAR_LENGTH(Namen.ZZeit) < 2,
              'keine Zeit vergeben',
              CONCAT(LEFT(Namen.ZZeit,2), ':', RIGHT(Namen.ZZeit,2), ' Uhr')
            )
          ) AS Zeit
        FROM
          Namen
        LEFT JOIN
          Buchungen
        ON
          Namen.Jahr = Buchungen.Jahr
        AND
          Namen.Bnr = Buchungen.Bnr
        LEFT JOIN
          Zustiege
        ON
          Namen.ZOrt = Zustiege.Kurz
        WHERE
          Buchungen.Storno IS NULL
        AND
          Zustiege.Preis = 0
        AND
          Namen.Jahr = """ + jahr + """
        AND
          Namen.Rnr = """ + rnr + """
        ORDER BY
          Namen.ZZeit,
          Namen.ZOrt,
          Namen.Bnr,
          Namen.PaxNr
        """

      cursor.execute(anfrage)
      uebergabe = cursor.fetchall()

      uebergabeliste = [
        {
          "id": i,
          "anrede": ub[0],
          "vorname": ub[1],
          "name": ub[2],
          "option": ub[3],
          "warteliste": ub[4],
          "selbebnr": True if i > 0 and ub[5] == uebergabe[i-1][5] else False,
          "bnr": ub[5],
          "selberort": True if i > 0 and ub[6] == uebergabe[i-1][6] else False,
          "ortkurz": ub[6],
          "ort": ub[7],
          "zeit": ub[8],
        }
        for i, ub in enumerate(uebergabe)
      ]


    finally:
      conn.close()

    locale.setlocale(locale.LC_TIME, "de_DE")

    document = Document()

    # Metadata
    core_properties = document.core_properties
    core_properties.author = 'Andreas Finger'
    core_properties.language = 'Deutsch'
    core_properties.category = 'Reiseablauf'
    core_properties.comments = 'zur Orientierung fuer Busfahrer'
    core_properties.content_status = 'Draft'
    #core_properties.created = date.today()
    #core_properties.created = ("%s" % datetime.strftime(datetime.now(), "%B"))
    core_properties.identifier = jahr + u" / " + rnr
    core_properties.keywords = "Reiseablauf"
    core_properties.last_modified_by = "af"
#    core_properties.title = re.sub(r'\xb7|\xdf|\xf6|\u2019', '-', reise.titel)
    core_properties.title = re.sub(r'[^a-zA-Z0-9_ ."-]', '-', docx_title)
#    core_properties.title = reiseziel

    #docx_title="reise_" + str(reiseziel) + ".docx"
    docx_title="reise_" + core_properties.title + ".docx"
    #document.add_picture('/var/www/reiseservice-schwerin/rss2' + static('images/logo_google_blau.png'), width=Inches(2))
    #sections = document.sections
    #sections[0].orientation = WD_ORIENT.LANDSCAPE
    #document.sections[0].orientation = WD_ORIENT.LANDSCAPE

    # Formatierungen
    font = document.styles['Normal'].font
    font.name = 'Submariner R24'
    font.size = Pt(12)
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1

    # Dokumentstruktur
    document.add_heading(u'Ablauf R' + rnr + u'/' + jahr + ': ' + ziel + ' vom ' + vom + ' bis ' + bis)
    document.add_paragraph()
    document.add_paragraph('1. Tag: ' + vom).bold = True

    p = document.add_paragraph()
    run = p.add_run()
    run.add_break()
    run.add_text('Zustiege:')
    run.add_break()

    #table = document.add_table(rows=len(uebergabeliste), cols=6)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Bnr'
    hdr_cells[1].text = 'Anrede'
    hdr_cells[2].text = 'Vorname'
    hdr_cells[3].text = 'Name'

    for i, ub in enumerate(uebergabeliste):
      if (i == 0) or (i > 0 and ub['ort'] != uebergabeliste[i-1]['ort']):
        row_cells = table.add_row().cells
        row_cells = table.add_row().cells
        row_cells[0].text = ub['zeit']
        row_cells[1].text = ub['ort']
        row_cells = table.add_row().cells
        row_cells = table.add_row().cells
        row_cells[0].text = 'B' + str(ub['bnr'])
        row_cells[1].text = ub['anrede']
        row_cells[2].text = ub['vorname']
        row_cells[3].text = ub['name']
      else:
        row_cells = table.add_row().cells
        row_cells[0].text = 'B' + str(ub['bnr'])
        row_cells[1].text = ub['anrede']
        row_cells[2].text = ub['vorname']
        row_cells[3].text = ub['name']

    p = document.add_paragraph()
    run = p.add_run()
    run.add_break()

    for ub in uebergabeliste:
      run.add_text('B' + str(ub['bnr']))
      run.add_break()
      run.add_text(ub['anrede'] + ' ')
      run.add_text(ub['vorname'] + ' ')
      run.add_text(ub['name'])
      run.add_break()
      run.add_text(ub['ort'])
      run.add_break()
      run.add_text(ub['zeit'])
      run.add_break()
      run.add_break()

    # -----------------------------
    # Prepare document for download
    # -----------------------------
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
      f.getvalue(),
      content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    response["Content-Encoding"] = "UTF-8"

    return response
